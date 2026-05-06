# -*- coding: utf-8 -*-
"""검증: 사업계획서 docx 5종 — 회사명/사업자번호/대표자/매출/4lang/차별점/마일스톤"""
import os, sys, io, re, json
from pathlib import Path
from docx import Document

ROOT = Path(__file__).parent
DOCS = [
    "사업계획서_AI리그_2026.docx",
    "사업계획서_AI리그_백업_재도전성공패키지_2026.docx",
    "사업계획서_KGlobal_2026.docx",
    "사업계획서_관광AI_2026.docx",
    "사업계획서_모두의창업_2026.docx",
]

CHECKS = {
    "회사명(쿤스튜디오)": [r"쿤스튜디오", r"KunStudio", r"Kun\s*Studio"],
    "사업자번호(552-59-00848)": [r"552-?59-?00848", r"552\s*-\s*59\s*-\s*00848"],
    "대표자(홍덕훈)": [r"홍덕훈", r"Hong\s*Dukhun", r"Deokhun"],
    "매출/매출0/잠재": [r"매출", r"₩\s*0", r"잠재", r"unlock", r"BEP", r"손익"],
    "자동화/schtask": [r"schtask", r"자동화", r"automation", r"스케줄"],
    "4언어(ko/en/ja/zh)": [r"ko\W*en\W*ja\W*zh", r"한국어.{0,10}영어.{0,10}일본어.{0,10}중국어", r"4\s*개\s*(언어|국어|국가)", r"글로벌\s*4"],
    "차별점(점신/포스텔러)": [r"점신", r"포스텔러", r"forceteller", r"competitor", r"차별"],
    "정통명리(음양/오행/십신)": [r"음양|오행|십신|12운성|신살|일주|월주"],
    "AI Q&A/매직링크": [r"AI\s*Q&A|AI\s*챗|챗봇|매직\s*링크|magic\s*link"],
    "마일스톤(5/20→6→9→12)": [r"5\s*/\s*20|마감|6월|9월|12월|마일스톤|로드맵"],
    "VC/투자(Antler/Kakao/D2SF)": [r"Antler|Kakao\s*Ventures|D2SF|VC|투자"],
}

def extract_text(p):
    d = Document(str(p))
    chunks = []
    for para in d.paragraphs:
        if para.text:
            chunks.append(para.text)
    for tbl in d.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if cell.text:
                    chunks.append(cell.text)
    return "\n".join(chunks)

def check(text):
    out = {}
    for label, patterns in CHECKS.items():
        hit = False
        snip = ""
        for pat in patterns:
            m = re.search(pat, text, re.IGNORECASE)
            if m:
                hit = True
                start = max(0, m.start() - 30)
                end = min(len(text), m.end() + 30)
                snip = text[start:end].replace("\n", " ")
                break
        out[label] = (hit, snip)
    return out

def main():
    report = []
    summary_table = []
    summary_table.append("| 파일 | " + " | ".join(CHECKS.keys()) + " |")
    summary_table.append("|---|" + "|".join(["---"] * len(CHECKS)) + "|")

    for name in DOCS:
        p = ROOT / name
        if not p.exists():
            report.append(f"### {name}\n\n**MISSING FILE**\n")
            continue
        try:
            text = extract_text(p)
        except Exception as e:
            report.append(f"### {name}\n\nERROR: {e}\n")
            continue
        result = check(text)
        word_count = len(text)
        para_count = len(text.split("\n"))
        row = [name]
        for label in CHECKS:
            hit, _ = result[label]
            row.append("OK" if hit else "MISS")
        summary_table.append("| " + " | ".join(row) + " |")

        report.append(f"### {name}\n")
        report.append(f"- 글자수: {word_count}, 단락수: {para_count}\n")
        for label, (hit, snip) in result.items():
            mark = "[OK]" if hit else "[MISS]"
            report.append(f"  - {mark} {label}: `{snip[:80]}`")
        report.append("")

    out = ["# 사업계획서 docx 5종 검증 보고 (2026-05-07)\n",
           "## 요약 매트릭스\n",
           "\n".join(summary_table),
           "\n\n## 상세\n",
           "\n".join(report)]
    out_path = Path("C:/Users/hdh02/Desktop/cheonmyeongdang/docs/grant_docx_validation_report.md")
    out_path.write_text("\n".join(out), encoding="utf-8")
    sys.stdout.buffer.write(("\n".join(summary_table) + "\n").encode("utf-8"))
    print(f"\nWritten: {out_path}")

if __name__ == "__main__":
    main()
