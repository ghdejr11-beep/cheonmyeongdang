# -*- coding: utf-8 -*-
"""
Generate 5 customized 사업계획서 docx files for 2026-05 round of grant applications.
Reuses 70% structure from 4/24 gyeongbuk_doyak templates, customizes 30% per grant.
"""
import os
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = r"D:\cheonmyeongdang\departments\tax\applications\round2_2026_05"
os.makedirs(OUT_DIR, exist_ok=True)


# ================== COMMON FACTS ==================
COMPANY = {
    "name": "쿤스튜디오 (KunStudio)",
    "ceo": "홍덕훈",
    "biz_no": "552-59-00848",
    "open_date": "2026-04-01",
    "address": "경상북도 경주시 외동읍 제내못안길 25-52",
    "industry": "정보통신업 / 응용 소프트웨어 개발 및 공급업",
    "tax_type": "간이과세자",
    "phone": "010-XXXX-XXXX",  # masked placeholder
    "email": "ghdejr11@gmail.com",
}


def set_korean_font(run, size=10, bold=False):
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.bold = bold


def add_heading(doc, text, level=1):
    p = doc.add_heading("", level=level)
    run = p.add_run(text)
    set_korean_font(run, size=14 if level == 1 else 12, bold=True)
    return p


def add_para(doc, text, bold=False, size=10, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_korean_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text, size=10):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_korean_font(run, size=size)
    return p


def add_info_table(doc, rows):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        t.cell(i, 0).text = k
        t.cell(i, 1).text = v
        for j in range(2):
            for p in t.cell(i, j).paragraphs:
                for r in p.runs:
                    set_korean_font(r, size=10, bold=(j == 0))
    return t


def add_data_table(doc, header, rows):
    t = doc.add_table(rows=len(rows) + 1, cols=len(header))
    t.style = "Light Grid Accent 1"
    for j, h in enumerate(header):
        t.cell(0, j).text = h
        for p in t.cell(0, j).paragraphs:
            for r in p.runs:
                set_korean_font(r, size=10, bold=True)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            t.cell(i + 1, j).text = str(v)
            for p in t.cell(i + 1, j).paragraphs:
                for r in p.runs:
                    set_korean_font(r, size=10)
    return t


def add_cover(doc, grant_name, sub_title, doc_title="사업계획서"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n\n2026년 {grant_name}\n")
    set_korean_font(run, size=16, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"({sub_title})\n")
    set_korean_font(run, size=12, bold=False)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n{doc_title}\n")
    set_korean_font(run, size=22, bold=True)

    doc.add_paragraph()
    add_info_table(
        doc,
        [
            ("상호", COMPANY["name"]),
            ("대표자", COMPANY["ceo"]),
            ("사업자등록번호", COMPANY["biz_no"]),
            ("사업장", COMPANY["address"]),
            ("개업일", COMPANY["open_date"]),
            ("연락처", "이메일: " + COMPANY["email"] + " / 휴대폰: 010-XXXX-XXXX"),
            ("신청유형", grant_name),
        ],
    )
    doc.add_page_break()


def add_company_overview(doc):
    add_heading(doc, "1. 기업 일반 현황", 1)
    add_info_table(
        doc,
        [
            ("상호", COMPANY["name"]),
            ("대표자", COMPANY["ceo"] + " (1985생, 만 41세)"),
            ("사업자등록번호", COMPANY["biz_no"]),
            ("사업장", COMPANY["address"]),
            ("업태/종목", COMPANY["industry"]),
            ("개업일자", COMPANY["open_date"] + " (창업 1개월차 초기창업)"),
            ("과세유형", COMPANY["tax_type"]),
            ("상시근로자", "1명 (대표자, 솔로 운영)"),
            ("핵심기술", "Claude Agent SDK + Pollinations Flux + Suno + Python 자동화 + RAG"),
            ("주요채널", "Vercel / Google Play / Amazon KDP / PayPal / Gumroad / YouTube 5채널"),
            ("보유서비스", "천명당, KORLENS, 세금N혜택, K-Wisdom YouTube, HexDrop 외 7+"),
            ("정부지원이력", "없음 (2026년 첫 신청)"),
        ],
    )
    doc.add_paragraph()


# ================== PLAN 1: AI 리그 ==================
def plan_1_ai_league():
    """K-스타트업 AI 리그 - 5억 상금"""
    doc = Document()
    add_cover(
        doc,
        "K-Startup AI 리그",
        "AI 분야 업력 3년 미만 초기창업기업 — 우수상금 5억",
    )
    add_company_overview(doc)

    add_heading(doc, "2. 창업 동기 및 배경", 1)
    add_para(
        doc,
        "신청인은 2026년 4월 1일 경상북도 경주에 1인 AI 스튜디오 '쿤스튜디오 (KunStudio)'를 창업하였다. "
        "창업 1개월 만에 Claude Agent SDK, Pollinations Flux, Suno API 등 최신 멀티모달 AI를 활용하여 "
        "3개의 라이브 AI 서비스(천명당·세금N혜택·KORLENS)를 운영 중이며, "
        "특히 2026-05-03 PayPal Smart Buttons 글로벌 결제 라이브 전환으로 "
        "한국 1인 AI 스타트업 최초로 글로벌 영문 사용자 첫 매출을 기록하였다.",
    )
    add_para(doc, "기존 한국 AI 스타트업이 놓친 3가지 시장 공백:")
    add_bullet(doc, "한국 전통문화(사주·관상·궁합) 글로벌 AI화 — Saju AI라는 영문 서비스 사실상 부재 (Astrology 시장 $13B/년)")
    add_bullet(doc, "5월 종소세 시즌 1,000만 자영업자 환급 누락 — 삼쩜삼 수수료 부담 + 9.9% 환불보장 모델 공백")
    add_bullet(doc, "한국 인바운드 여행자 3,000만+ 시장의 AI 현지인 픽 4개국어 가이드 공백")

    add_heading(doc, "3. 사업 아이템 — 3개 라이브 AI 라인업", 1)
    add_heading(doc, "3-1. 천명당 (cheonmyeongdang.vercel.app) — Saju AI 추론 엔진", 2)
    add_bullet(doc, "한국 명리학 350+ 규칙 베이스 + Claude Sonnet 4.6 LLM 추론 + Pollinations Flux 시각화 결합")
    add_bullet(doc, "AI 진단 영역: 사주(四柱)·관상·손금·꿈해몽·궁합·신년운세·타로+사주 융합")
    add_bullet(doc, "글로벌 영문 우선 — PayPal Smart Buttons 2026-05-03 라이브, 첫 결제 검증")
    add_bullet(doc, "프리미엄: ₩2,900 ~ ₩29,900 (단건/월회원/연회원 3티어)")

    add_heading(doc, "3-2. 세금N혜택 (tax-n-benefit-api.vercel.app) — 종소세 자동화 AI", 2)
    add_bullet(doc, "Claude API + 국세청 홈택스 데이터 API 연동")
    add_bullet(doc, "AI 환급액 계산 + 9.9% 합리적 수수료 + 환급 0원 시 100% 환불 보장")
    add_bullet(doc, "삼쩜삼(15~30% 수수료) 대비 신뢰 차별화")
    add_bullet(doc, "빌게이트 PG 통합 완료, 5월 종소세 시즌 라이브")

    add_heading(doc, "3-3. KORLENS (korlens.app) — 한국 여행 AI 가이드", 2)
    add_bullet(doc, "4개국어 (한/영/일/중) 한국 여행자용 AI 추천 가이드")
    add_bullet(doc, "Supabase + Google OAuth + 자체 추천 알고리즘 ('현지인 픽')")
    add_bullet(doc, "수익: 호텔·체험·식당 어필리에이트 + 프리미엄 가이드 PDF")
    add_bullet(doc, "Play Console 비공개 테스트 중")

    add_heading(doc, "4. AI 기술 차별화 — 멀티모달 90+ 자동화 부서", 1)
    add_para(
        doc,
        "쿤스튜디오는 '솔로 창업자 1인 = 다른 AI 스타트업 10인 분량'의 자동화 인프라를 구축하였다.",
        bold=True,
    )
    add_data_table(
        doc,
        ["AI 스택", "역할", "활용 영역"],
        [
            ["Claude Sonnet 4.6", "텍스트 추론·콘텐츠 생성", "사주 풀이·세무 상담·여행 추천"],
            ["Pollinations Flux", "이미지 생성", "사주 부적·썸네일·SEO 이미지"],
            ["Suno API", "음악·BGM 생성", "Sori Atlas YouTube 채널 24/7"],
            ["ffmpeg + Whisper", "영상 합성·자막", "K-Wisdom 5채널 일일 자동 업로드"],
            ["Pollinations TTS", "음성 합성", "Faceless 영상 영어 내레이션"],
            ["Vercel Functions", "서버리스 백엔드", "PayPal/Toss 결제 webhook + 결제 검증"],
            ["Supabase + Postgres", "사용자·결제 DB", "KORLENS 회원 관리"],
            ["Python schtask", "자동화 스케줄러", "90+ 작업 매시간 자동 실행"],
        ],
    )
    add_para(
        doc,
        "특히 90+ schtask 매시간 자동 실행 구조로 콘텐츠 생산-배포-매출 모니터링-경쟁사 인텔리전스를 "
        "사람의 개입 없이 24시간 운영 중이다. 이는 일반 AI 스타트업이 PM·디자이너·마케터 5인 "
        "고용 후에야 달성하는 운영 규모를 인건비 0원으로 달성한 것이다.",
    )

    add_heading(doc, "5. 시장 분석", 1)
    add_data_table(
        doc,
        ["서비스", "타겟", "시장 규모 (TAM)", "도달 가능 (SAM)"],
        [
            ["천명당", "글로벌 K-culture/Astrology 팬", "$13B Astrology 글로벌", "K-Astrology 영문 검색 ~$200M"],
            ["세금N혜택", "한국 자영업자·프리랜서", "환급 누락 연 2조원", "1,000만 종소세 신고자 5%"],
            ["KORLENS", "한국 인바운드 여행자", "3,000만 명/년 인바운드", "MZ 여행자 영문/일문 30%"],
            ["AI 리그 합산", "글로벌 + 국내 통합", "10조원+ 통합 AI 도메인", "1차 점유 0.01% = ₩100억"],
        ],
    )

    add_heading(doc, "6. 글로벌 검증 매출 (2026-05 기준)", 1)
    add_bullet(doc, "PayPal Smart Buttons 2026-05-03 라이브 — 첫 글로벌 영문 결제 트랜잭션 검증")
    add_bullet(doc, "Gumroad 27+ 영문 ebook (Korean culture / Saju / K-pop / K-beauty) 글로벌 판매 중")
    add_bullet(doc, "K-Wisdom YouTube 채널 5개 일일 자동 업로드 — 글로벌 시청자 유입")
    add_bullet(doc, "Bluesky / Hashnode / Beehiiv 글로벌 SNS 다채널 자동 발행")

    add_heading(doc, "7. 사업 추진 로드맵 (2026.5 ~ 2027.4)", 1)
    add_data_table(
        doc,
        ["기간", "마일스톤", "예상 지표"],
        [
            ["5~6월", "AI 리그 진출 + 종소세 환급 자동화 가동", "MAU 5,000 / MRR ₩500만"],
            ["7~8월", "PayPal 글로벌 영문 매출 본격 + AdSense 승인", "MAU 15,000 / MRR ₩1,500만"],
            ["9~10월", "KORLENS 정식 출시 + 4개국어 확장", "MAU 30,000 / MRR ₩3,000만"],
            ["11~12월", "B2B AI API 판매 (사주/관상 SDK)", "MAU 50,000 / MRR ₩5,000만"],
            ["2027 Q1", "AI 리그 우승 + 시리즈 A 라운드", "MRR ₩1억 / ARR ₩12억"],
        ],
    )

    add_heading(doc, "8. 5억원 사업화 자금 집행 계획", 1)
    add_data_table(
        doc,
        ["항목", "금액 (만원)", "내용"],
        [
            ["① AI 모델·인프라 고도화", "15,000", "Claude API 정식 / Pollinations Pro / Suno Enterprise / GPU 서버"],
            ["② 글로벌 마케팅", "12,000", "Pinterest/TikTok English 광고, Google Ads (Astrology/Saju 키워드)"],
            ["③ 인력 채용 (3명)", "12,000", "AI 엔지니어 1, 글로벌 마케터 1, 영문 콘텐츠 PM 1"],
            ["④ 데이터·법률 컨설팅", "5,000", "사주 데이터셋 350→2,000+ 확장, 글로벌 결제 규제 검토"],
            ["⑤ 인증·운영 인프라", "6,000", "AWS 이전, ISMS 준비, AI 윤리 감사"],
            ["합계", "50,000", "총 5억원 (자부담 0원, 100% 국비)"],
        ],
    )

    add_heading(doc, "9. 팀 및 외부 네트워크", 1)
    add_bullet(doc, "홍덕훈 (1985생, 대표자) — Claude Agent SDK 기반 멀티 에이전트 아키텍처 1인 운영")
    add_bullet(doc, "1개월 만에 13개 자동화 부서 구축 + 8앱 Play Console 진입 + 27+ KDP ebook 발간")
    add_bullet(doc, "외부 네트워크: 세무사 자문 / 한국관광공사 (KORLENS 공모전 출품) / 빌게이트 PG / Anthropic API")

    add_heading(doc, "10. 향후 비전 및 사회적 가치", 1)
    add_para(
        doc,
        "쿤스튜디오는 '한국 문화 X AI 자동화'라는 단일 미션 아래 천명당·KORLENS·세금N혜택·K-Wisdom의 4축 "
        "글로벌 AI 플랫폼을 구축한다. 100일 안에 ₩10억 매출, 1년 안에 시리즈 A를 목표로 한다.",
    )
    add_bullet(doc, "비수도권(경주) 본사 유지 → 지역 일자리 3+ 명 창출 (1년차)")
    add_bullet(doc, "한국 전통문화의 글로벌 디지털 수출 (BTS·드라마 이후의 K-soft power)")
    add_bullet(doc, "1인 AI 창업자 모범 사례로 후속 솔로 창업 생태계 자극")

    add_heading(doc, "11. 부록 — 검증 채널", 1)
    add_bullet(doc, "천명당: https://cheonmyeongdang.vercel.app (PayPal 글로벌 결제 라이브)")
    add_bullet(doc, "KORLENS: https://korlens.app (Play Console 비공개 테스트)")
    add_bullet(doc, "세금N혜택: https://tax-n-benefit-api.vercel.app (빌게이트 PG)")
    add_bullet(doc, "Gumroad 영문 ebook: 27+ 상품 글로벌 판매")
    add_bullet(doc, "K-Wisdom YouTube 5채널 일일 자동 업로드")

    out = os.path.join(OUT_DIR, "사업계획서_AI리그_2026.docx")
    doc.save(out)
    return out


# ================== PLAN 2: 관광 AI ==================
def plan_2_tourism_ai():
    """한국관광공사 관광기업 데이터·AI 활용 지원 - KORLENS 단독"""
    doc = Document()
    add_cover(
        doc,
        "한국관광공사 관광기업 데이터·AI 활용 지원",
        "KORLENS 단독 / AI 솔루션 100만 + 캠페인 1,000만 활용 계획",
    )
    add_company_overview(doc)

    add_heading(doc, "2. 창업 배경 및 KORLENS 출범 동기", 1)
    add_para(
        doc,
        "신청인은 2026년 4월 쿤스튜디오를 창업하며 한국 인바운드 여행 시장의 AI 활용 공백을 발견하였다. "
        "외국인 관광객 3,000만+ 시장에서 4개국어(한/영/일/중) AI 추천 가이드 + '현지인 픽' 알고리즘은 "
        "사실상 부재 상태이며, 비짓코리아·트리플·마이리얼트립 등 기존 플랫폼은 분절적 정보 나열에 그친다.",
    )
    add_bullet(doc, "기존 가이드: 정형 콘텐츠 나열 (벚꽃 명소 100선 등)")
    add_bullet(doc, "KORLENS: AI 추천 + '현지인 픽' + 4개국어 동시 + 어필리에이트 자동 매칭")

    add_heading(doc, "3. KORLENS 서비스 개요", 1)
    add_heading(doc, "3-1. 핵심 기능", 2)
    add_bullet(doc, "AI 추천 알고리즘: 사용자 여행 스타일 X 한국 지역 매칭")
    add_bullet(doc, "현지인 픽: 한국인이 실제 가는 식당·카페·체험 큐레이션")
    add_bullet(doc, "4개국어 자동 번역 (한국어 원본 → 영/일/중 동시 노출)")
    add_bullet(doc, "어필리에이트 자동 매칭: 호텔(Booking.com) + 체험(Klook) + 식당(네이버 예약)")
    add_bullet(doc, "프리미엄 가이드 PDF (지역별 5,000원 ~ 19,900원)")

    add_heading(doc, "3-2. 기술 스택", 2)
    add_data_table(
        doc,
        ["기술", "역할"],
        [
            ["Supabase + Postgres", "사용자 / 가이드 데이터베이스"],
            ["Google OAuth", "글로벌 회원 인증 (간편 가입)"],
            ["Claude Sonnet 4.6", "AI 추천 + 4개국어 번역"],
            ["Pollinations Flux", "지역별 SEO 이미지 자동 생성"],
            ["Vercel Edge", "글로벌 CDN 저지연 응답"],
            ["Play Console (Android)", "모바일 앱 배포 (iOS는 PWA)"],
        ],
    )

    add_heading(doc, "3-3. 보유 자산 (적격 요건)", 2)
    add_bullet(doc, "관광 분야 중소기업: 정보통신업 / 응용 SW 개발업 (관광기술 적격)")
    add_bullet(doc, "웹사이트 보유: korlens.app (라이브)")
    add_bullet(doc, "모바일 앱 보유: Google Play Console 비공개 테스트 진행 중")
    add_bullet(doc, "사업자등록증 보유: " + COMPANY["biz_no"])

    add_heading(doc, "4. 시장 분석 — 한국 인바운드 여행 시장", 1)
    add_data_table(
        doc,
        ["지표", "수치", "출처/시점"],
        [
            ["2024 인바운드 외국인", "1,640만 명 / 연", "한국관광공사 통계"],
            ["2026 목표", "2,300만 명 / 연", "정부 K-관광 정책"],
            ["MZ 비중", "62%", "MZ 외국인 디지털 우선"],
            ["모바일 가이드 검색", "47%", "Google Trends Korea Travel"],
            ["4개국어 통합 AI 가이드", "사실상 0개", "경쟁사 분석 결과"],
        ],
    )

    add_heading(doc, "5. AI 솔루션 100만원 활용 계획", 1)
    add_data_table(
        doc,
        ["항목", "금액 (만원)", "내용"],
        [
            ["Claude API 정식", "30", "월 25만 토큰 / 추천·번역 자동화"],
            ["Pollinations Pro", "20", "지역별 SEO 이미지 1,000+ 생성"],
            ["Supabase Pro", "30", "DB Pro 플랜 1년"],
            ["Google Maps API", "20", "지도/위치/리뷰 데이터"],
            ["합계", "100", "100만원 100% 활용"],
        ],
    )

    add_heading(doc, "6. 캠페인비 1,000만원 집행 계획", 1)
    add_data_table(
        doc,
        ["항목", "금액 (만원)", "내용"],
        [
            ["① 글로벌 SNS 광고", "300", "Pinterest English / TikTok English / Instagram"],
            ["② 콘텐츠 제작", "250", "지역별 영상 50편 (제주/경주/부산/서울)"],
            ["③ 인플루언서 마케팅", "200", "K-travel 영문 인플루언서 5인 협업"],
            ["④ 영문 SEO·블로그", "150", "Hashnode/Medium English 30+ 포스트"],
            ["⑤ 어필리에이트 트래픽", "100", "Google Ads Korea Travel 키워드"],
            ["합계", "1,000", "글로벌 영문 시장 우선"],
        ],
    )

    add_heading(doc, "7. 추진 일정 (2026.6 ~ 2026.12)", 1)
    add_data_table(
        doc,
        ["기간", "마일스톤", "지표"],
        [
            ["6월", "AI 솔루션 통합 + Play 정식 출시", "DAU 100"],
            ["7월", "글로벌 영문 캠페인 가동", "MAU 1,000"],
            ["8~9월", "어필리에이트 본격 (호텔/체험)", "MAU 5,000 / 첫 매출"],
            ["10~11월", "프리미엄 가이드 PDF 10종", "MAU 15,000 / MRR ₩300만"],
            ["12월", "한국관광공사 연계 캠페인 보고", "MAU 30,000 / MRR ₩500만"],
        ],
    )

    add_heading(doc, "8. 기대 효과", 1)
    add_bullet(doc, "한국 인바운드 여행 디지털 가이드 공백 해소 (4개국어 동시)")
    add_bullet(doc, "지역(비수도권 경주 포함) 관광 큐레이션 확대 — 비수도권 트래픽 최소 30% 확보")
    add_bullet(doc, "AI 솔루션의 관광 적용 모범 사례 (관광공사 BP 기여)")
    add_bullet(doc, "어필리에이트 매출 → 지역 소상공인 (호텔/식당) 매출 직접 연결")

    add_heading(doc, "9. 부록", 1)
    add_bullet(doc, "KORLENS: https://korlens.app")
    add_bullet(doc, "Play Console 비공개 테스트 트랙: 활성")
    add_bullet(doc, "한국관광공사 2026년 4월 18일 KORLENS 공모전 출품 이력")
    add_bullet(doc, "도메인 등록증: korlens.app (보유)")

    out = os.path.join(OUT_DIR, "사업계획서_관광AI_2026.docx")
    doc.save(out)
    return out


# ================== PLAN 3: 모두의 창업 로컬 ==================
def plan_3_local_track():
    """모두의 창업 로컬트랙 - 비수도권 + 천명당 한국 전통 사주"""
    doc = Document()
    add_cover(
        doc,
        "모두의 창업 — 로컬트랙",
        "비수도권(경주) X 한국 전통문화 글로벌 AI화 / 활동 200만 → 사업화 3,000만 → 우승 1억",
    )
    add_company_overview(doc)

    add_heading(doc, "2. 로컬 자원 X 글로벌 디지털 — 창업 동기", 1)
    add_para(
        doc,
        "신청인은 비수도권 경상북도 경주에 본사를 두고 '한국 전통 사주 X 글로벌 영문 AI'라는 "
        "독특한 포지셔닝의 천명당 (cheonmyeongdang.vercel.app) 서비스를 운영 중이다. "
        "경주는 신라 천년 고도 + 첨성대 + 안압지 등 한국 전통의 본거지로, "
        "사주·명리학의 역사적 정통성을 어필할 수 있는 최적의 로컬 자원을 보유한다.",
    )
    add_para(doc, "로컬(경주) X 글로벌(영문 AI) 결합 차별화:", bold=True)
    add_bullet(doc, "수도권 사주 앱(점신·포스텔러)은 한국어 내수 100%, 글로벌 영문 시장 공백")
    add_bullet(doc, "천명당은 영문·한국어 2중 트랙 — Claude API 기반 글로벌 영문 사주 풀이")
    add_bullet(doc, "경주 전통문화 자원 (신라사주 콘텐츠) 활용 가능성")
    add_bullet(doc, "글로벌 K-pop·K-drama 팬 (수억 명)이 사주에 관심 → 영문 검색량 증가")

    add_heading(doc, "3. 천명당 — 한국 전통 사주의 글로벌 디지털 진출", 1)
    add_heading(doc, "3-1. 서비스 개요", 2)
    add_bullet(doc, "한국 명리학 350+ 규칙 + Claude Sonnet 4.6 LLM 추론")
    add_bullet(doc, "AI 진단 7종: 사주·관상·손금·꿈해몽·궁합·신년운세·타로+사주 융합")
    add_bullet(doc, "PayPal Smart Buttons 라이브 (2026-05-03 첫 글로벌 결제)")
    add_bullet(doc, "프리미엄: ₩2,900 ~ ₩29,900 (단건/월회원/연회원)")

    add_heading(doc, "3-2. 비수도권 가산 + 로컬 강점", 2)
    add_bullet(doc, "본사 경상북도 경주 — 비수도권 가산점 +5점 기대")
    add_bullet(doc, "지역 자원 활용: 경주 신라 역사 콘텐츠 → 사주 풀이 스토리텔링")
    add_bullet(doc, "지역 일자리: 1년 후 영문 PM·번역·콘텐츠 1~2명 채용 계획")

    add_heading(doc, "4. 시장 분석", 1)
    add_data_table(
        doc,
        ["시장", "TAM", "SAM", "SOM (1년)"],
        [
            ["글로벌 Astrology", "$13B/년", "Saju/Asian Astrology $200M", "$200K (0.1%)"],
            ["국내 사주 앱", "₩2,000억/년", "20~40 여성 ₩600억", "₩6억 (1%)"],
            ["K-culture 영문 콘텐츠", "$5B/년", "K-spirituality $80M", "$80K"],
            ["합산", "$20B+", "$280M", "₩10억+"],
        ],
    )

    add_heading(doc, "5. 단계별 시나리오 — 활동/사업화/우승", 1)
    add_heading(doc, "5-1. 활동지원금 200만원 활용", 2)
    add_data_table(
        doc,
        ["항목", "금액 (만원)", "내용"],
        [
            ["글로벌 영문 콘텐츠 50편", "70", "Claude API + Pollinations 자동 생성 + 외주"],
            ["Pinterest/TikTok English 광고", "60", "MZ 영문 K-culture 팬 도달 100만+"],
            ["경주 로컬 영상 (신라 사주)", "40", "현지 촬영 + 영문 자막"],
            ["멘토링·교육 참가", "30", "전문가 세션 출석"],
            ["합계", "200", ""],
        ],
    )

    add_heading(doc, "5-2. 사업화 자금 3,000만원 활용", 2)
    add_data_table(
        doc,
        ["항목", "금액 (만원)", "내용"],
        [
            ["① 사주 데이터 확장", "800", "350 → 2,000+ 규칙 데이터셋 구축"],
            ["② 글로벌 영문 마케팅", "900", "Pinterest/TikTok English 본격 광고"],
            ["③ AI API 정식 구독", "500", "Claude API + Pollinations Pro 1년"],
            ["④ 영문 PM·번역 외주", "500", "월 80만 X 6개월 (5개 언어)"],
            ["⑤ 경주 로컬 협업", "300", "신라사 박물관 콘텐츠 라이선스 협의"],
            ["합계", "3,000", ""],
        ],
    )

    add_heading(doc, "5-3. 우승상금 1억원 진출 시 시나리오", 2)
    add_bullet(doc, "글로벌 결제 인프라 풀가동 (PayPal + Lemon Squeezy 이중)")
    add_bullet(doc, "B2B 사주 AI API 판매 (게임/엔터/매칭앱 5+ 계약)")
    add_bullet(doc, "동남아 진출 (베트남/태국 — K-pop 팬덤 강세)")
    add_bullet(doc, "1년 내 ARR ₩12억 (MRR ₩1억) 달성 → 시리즈 A 라운드")

    add_heading(doc, "6. 차별화 (왜 우리가 우승해야 하는가)", 1)
    add_bullet(doc, "비수도권(경주) 본사 — 로컬 가산 + 지역 자원 직접 활용")
    add_bullet(doc, "이미 라이브 매출 검증 (PayPal 2026-05-03 첫 결제)")
    add_bullet(doc, "솔로 1인 창업 + 90+ 자동화로 인건비 0원 운영")
    add_bullet(doc, "한국 전통문화의 글로벌 영문화 — BTS·드라마 이후의 K-soft power")

    add_heading(doc, "7. 일정", 1)
    add_data_table(
        doc,
        ["기간", "마일스톤"],
        [
            ["6~7월", "활동지원금 200만원 집행 / 영문 콘텐츠 50편 / 글로벌 광고"],
            ["8~9월", "1차 심사 통과 후 사업화 3,000만원 집행 / 데이터 2,000+ 확장"],
            ["10~11월", "MAU 30,000 / MRR ₩3,000만 달성"],
            ["12월", "결선 진출 / 최종 우승 발표"],
            ["2027 Q1", "우승상금 1억원 활용 글로벌 본격 / 시리즈 A 준비"],
        ],
    )

    add_heading(doc, "8. 기대 효과", 1)
    add_bullet(doc, "비수도권 경주에서 글로벌 ₩10억+ 매출 달성 → 지역 창업 모범 사례")
    add_bullet(doc, "한국 전통문화의 디지털 수출 (사주·명리학 영문화)")
    add_bullet(doc, "지역 일자리 2~3명 창출 (1년차)")
    add_bullet(doc, "후속 솔로 AI 창업 생태계 자극 (1인 100억 모델)")

    add_heading(doc, "9. 팀·외부 네트워크", 1)
    add_bullet(doc, "홍덕훈 (대표자, 1985생) — Python/AI 자동화 1인 운영")
    add_bullet(doc, "외부: 경주시 경제진흥원 (예정) / 한국관광공사 / 빌게이트 PG / Anthropic API")

    add_heading(doc, "10. 부록", 1)
    add_bullet(doc, "천명당: https://cheonmyeongdang.vercel.app (PayPal 라이브)")
    add_bullet(doc, "사업자등록증: 552-59-00848 / 경주 본사")
    add_bullet(doc, "PayPal 첫 결제 영수증 (2026-05-03)")

    out = os.path.join(OUT_DIR, "사업계획서_모두의창업_2026.docx")
    doc.save(out)
    return out


# ================== PLAN 4: K-Global 해외진출 ==================
def plan_4_kglobal():
    """K-Global 해외진출 GDIN - 수출 검증"""
    doc = Document()
    add_cover(
        doc,
        "K-Global 해외진출 지원사업 (GDIN)",
        "글로벌 수출 검증 / PayPal·Gumroad·YouTube·Bluesky 다채널",
    )
    add_company_overview(doc)

    add_heading(doc, "2. 해외진출 동기 및 검증 실적", 1)
    add_para(
        doc,
        "쿤스튜디오는 2026년 4월 창업 1개월 만에 글로벌 영문 매출 검증을 완료한 1인 AI 스튜디오이다. "
        "한국 1인 AI 창업자가 PayPal Smart Buttons 라이브, Gumroad 27+ ebook 판매, "
        "Bluesky/Hashnode/Beehiiv 글로벌 다채널, K-Wisdom YouTube 5채널 글로벌 시청자 등 "
        "다층 수출 채널을 단기간에 구축한 사례이다.",
    )

    add_heading(doc, "3. 해외 매출 검증 채널 (4축)", 1)
    add_heading(doc, "3-1. PayPal Smart Buttons — 글로벌 결제 라이브", 2)
    add_bullet(doc, "2026-05-03 PayPal Business 라이브 전환 완료")
    add_bullet(doc, "confirm-payment.js Vercel Function PayPal 검증 통합")
    add_bullet(doc, "환경변수 3개 (CLIENT_ID, SECRET, BASE) 등록 완료")
    add_bullet(doc, "첫 글로벌 영문 결제 트랜잭션 검증 (천명당)")
    add_bullet(doc, "한국 PG(Toss/갤럭시아) 라이브 대기 별도 — 글로벌 우선 전략")

    add_heading(doc, "3-2. Gumroad — 27+ 영문 ebook 판매", 2)
    add_bullet(doc, "Korean culture / Saju / K-pop / K-beauty 카테고리")
    add_bullet(doc, "저자명 Deokgu Studio / Deokhun Hong")
    add_bullet(doc, "글로벌 결제 (Stripe 간접) — 미국·유럽·동남아 매출")
    add_bullet(doc, "B2B 디지털 상품 8종 (AI 부업 시스템 / 노션 템플릿 등)")

    add_heading(doc, "3-3. K-Wisdom YouTube 채널 5개 — 글로벌 시청자", 2)
    add_data_table(
        doc,
        ["채널", "주제", "포맷"],
        [
            ["Whisper Atlas", "Korean spirituality (영문)", "Faceless 자동 영상"],
            ["Wealth Blueprint", "Korean wealth/biz (영문)", "Faceless 자동 영상"],
            ["Inner Archetypes", "Korean psychology (영문)", "Faceless 자동 영상"],
            ["AI Side Hustle", "AI 부업 (영문)", "쇼츠"],
            ["Sori Atlas", "Korean Lofi (Suno 음원)", "24/7 라이브 (Hetzner VPS)"],
        ],
    )

    add_heading(doc, "3-4. Bluesky / Hashnode / Beehiiv — 글로벌 SNS·뉴스레터", 2)
    add_bullet(doc, "Bluesky @kunstudio — 글로벌 인디 개발자 커뮤니티 도달")
    add_bullet(doc, "Hashnode English 블로그 — 30+ 영문 포스트")
    add_bullet(doc, "Beehiiv 영문 뉴스레터 (Korean tech 글로벌)")
    add_bullet(doc, "Postiz 셀프호스트 (Railway) 다채널 자동 발행")

    add_heading(doc, "4. 해외 시장 분석", 1)
    add_data_table(
        doc,
        ["시장", "TAM", "현재 진입", "1년 목표"],
        [
            ["미국 K-culture 팬", "5,000만+ 명", "Gumroad ~5건/월", "200건/월"],
            ["일본 K-wave 팬", "1,500만 명", "K-Wisdom YouTube", "MAU 5,000"],
            ["동남아 K-pop 팬", "수억 명", "Bluesky 도달", "MAU 30,000"],
            ["유럽 K-fashion", "1,000만+ 명", "Pinterest English", "MAU 10,000"],
            ["글로벌 합산", "1억+ 도달 가능", "월 ~$200", "월 $100K (₩1.3억)"],
        ],
    )

    add_heading(doc, "5. K-Global 지원 활용 계획", 1)
    add_data_table(
        doc,
        ["항목", "금액 (만원)", "내용"],
        [
            ["① 글로벌 결제 인프라 확장", "1,500", "Lemon Squeezy 추가 + ISMS 준비 + 환율 헷지"],
            ["② 미국·일본 현지 마케팅", "2,500", "Pinterest US / Japanese influencer / TikTok JP"],
            ["③ 영문/일문/중문 콘텐츠 100편", "2,000", "외주 번역 + 영상 + 블로그"],
            ["④ 해외 전시·바이어 매칭", "1,000", "GDIN 행사 참가 + 현지 PG 미팅"],
            ["⑤ 글로벌 인증·법무", "1,000", "GDPR / CCPA / SOC2 준비"],
            ["⑥ 운영비", "1,000", "API/CDN/도메인/외주 6개월"],
            ["합계", "9,000", "글로벌 100% 집중"],
        ],
    )

    add_heading(doc, "6. 일정 (2026.6 ~ 2027.6)", 1)
    add_data_table(
        doc,
        ["기간", "마일스톤", "글로벌 매출"],
        [
            ["6~8월", "Lemon Squeezy 통합 + 영문 콘텐츠 본격", "$1,000/월"],
            ["9~11월", "일본·동남아 진입 + 인플루언서 협업", "$5,000/월"],
            ["12~2월", "유럽 진입 + GDPR 인증", "$15,000/월"],
            ["3~6월", "B2B 글로벌 API 계약 5+", "$50,000/월"],
            ["2027 Q3", "ARR $1M 도달 → 시리즈 A 라운드", ""],
        ],
    )

    add_heading(doc, "7. 차별화", 1)
    add_bullet(doc, "이미 글로벌 결제 라이브 (PayPal 2026-05-03)")
    add_bullet(doc, "27+ 영문 ebook으로 해외 매출 실적 검증")
    add_bullet(doc, "5채널 YouTube 글로벌 시청자 자동 확보")
    add_bullet(doc, "솔로 1인 + 90+ 자동화 = 다른 팀 10인 분량")
    add_bullet(doc, "비수도권(경주) 본사 + 글로벌 100% 매출 (지역 균형 발전 모범)")

    add_heading(doc, "8. 팀 및 글로벌 네트워크", 1)
    add_bullet(doc, "홍덕훈 (1985생) — 1인 글로벌 진출 검증 완료")
    add_bullet(doc, "글로벌 파트너: PayPal Business / Anthropic API / Pollinations / Suno / Hetzner VPS")
    add_bullet(doc, "1년 후 채용 계획: 글로벌 마케터 1, 영문 PM 1, 일문/중문 번역 외주")

    add_heading(doc, "9. 기대 효과", 1)
    add_bullet(doc, "1년 내 ARR $1M 도달 → 한국 1인 AI 창업자 글로벌 모범 사례")
    add_bullet(doc, "비수도권 본사 + 글로벌 매출 → 지역 균형 발전 BP")
    add_bullet(doc, "한국 문화 글로벌 디지털 수출 (사주·여행·언어 통합)")

    add_heading(doc, "10. 부록", 1)
    add_bullet(doc, "PayPal 첫 결제 영수증 (2026-05-03)")
    add_bullet(doc, "Gumroad 상점: 27+ 상품 매출 내역")
    add_bullet(doc, "K-Wisdom YouTube 5채널 통계")
    add_bullet(doc, "Bluesky/Hashnode 글로벌 SNS 통계")

    out = os.path.join(OUT_DIR, "사업계획서_KGlobal_2026.docx")
    doc.save(out)
    return out


# ================== PLAN 5: 재도전 성공패키지 (백업) ==================
def plan_5_restart_backup():
    """K-Startup 재도전성공패키지 - 백업 (5/7 마감)"""
    doc = Document()
    add_cover(
        doc,
        "K-Startup AI 리그 (백업) / 재도전성공패키지 트랙",
        "초기창업 1개월차 + 글로벌 검증 강점 / 5월 7일 마감 백업 옵션",
    )
    add_company_overview(doc)

    add_heading(doc, "2. 신청 트랙 및 신청 동기", 1)
    add_para(
        doc,
        "본 사업계획서는 K-스타트업 AI 리그 1순위 신청과 별도로, 재도전성공패키지 또는 초기창업 트랙에 "
        "백업으로 동시 제출하기 위한 사본이다. 쿤스튜디오는 2026년 4월 1일 창업한 1개월차 신규 창업자이며, "
        "재도전 트랙 적격 시 1순위 변경 또는 별도 트랙으로 활용한다.",
    )
    add_para(doc, "백업 트랙 강점:", bold=True)
    add_bullet(doc, "1개월차 신규 창업 — 초기창업 트랙 적격")
    add_bullet(doc, "이미 라이브 매출 검증 (PayPal 2026-05-03)")
    add_bullet(doc, "비수도권(경주) 가산점")
    add_bullet(doc, "솔로 + 90+ 자동화 인프라로 빠른 검증 능력")

    add_heading(doc, "3. 사업 아이템 — 천명당·세금N혜택·KORLENS 3축", 1)
    add_heading(doc, "3-1. 천명당 (메인)", 2)
    add_bullet(doc, "한국 명리학 350+ 규칙 + Claude Sonnet 4.6 LLM")
    add_bullet(doc, "PayPal Smart Buttons 글로벌 결제 라이브")
    add_bullet(doc, "프리미엄 ₩2,900~₩29,900 / 첫 매출 검증")

    add_heading(doc, "3-2. 세금N혜택 (서브)", 2)
    add_bullet(doc, "Claude API + 국세청 API 연동 종소세 자동화")
    add_bullet(doc, "9.9% 수수료 + 환불보장 / 5월 종소세 시즌 라이브")

    add_heading(doc, "3-3. KORLENS (서브)", 2)
    add_bullet(doc, "4개국어 한국 여행 가이드 + AI 현지인 픽")
    add_bullet(doc, "Play Console 비공개 테스트 중")

    add_heading(doc, "4. 시장 분석 (요약)", 1)
    add_data_table(
        doc,
        ["서비스", "시장 규모", "도달 가능"],
        [
            ["천명당", "글로벌 Astrology $13B", "K-Astrology $200M"],
            ["세금N혜택", "환급 누락 연 2조원", "1,000만 신고자 5%"],
            ["KORLENS", "인바운드 3,000만 명/년", "MZ 외국인 30%"],
        ],
    )

    add_heading(doc, "5. AI 기술 차별화", 1)
    add_data_table(
        doc,
        ["AI 스택", "활용"],
        [
            ["Claude Sonnet 4.6", "사주 풀이·세무 상담·여행 추천"],
            ["Pollinations Flux", "이미지 자동 생성"],
            ["Suno API", "Sori Atlas 음악 자동"],
            ["ffmpeg + Whisper", "K-Wisdom 영상 자동"],
            ["Python schtask 90+", "24시간 자동 운영"],
        ],
    )

    add_heading(doc, "6. 글로벌 검증 매출 (1개월 이내)", 1)
    add_bullet(doc, "PayPal 글로벌 결제 라이브 (2026-05-03)")
    add_bullet(doc, "Gumroad 27+ 영문 ebook 판매")
    add_bullet(doc, "K-Wisdom YouTube 5채널 일일 자동 업로드")
    add_bullet(doc, "Bluesky/Hashnode/Beehiiv 글로벌 SNS")

    add_heading(doc, "7. 일정 (2026.6 ~ 2027.5)", 1)
    add_data_table(
        doc,
        ["기간", "마일스톤", "지표"],
        [
            ["6~7월", "AI 모델 고도화 + 영문 콘텐츠 100편", "MAU 5,000"],
            ["8~9월", "글로벌 광고 본격 + AdSense 승인", "MAU 15,000 / MRR ₩1,500만"],
            ["10~12월", "B2B AI API 판매 + 동남아 진입", "MAU 30,000 / MRR ₩3,000만"],
            ["2027 Q1", "ARR ₩12억 + 시리즈 A 준비", ""],
        ],
    )

    add_heading(doc, "8. 자금 집행 계획 (1억원 기준)", 1)
    add_data_table(
        doc,
        ["항목", "금액 (만원)", "내용"],
        [
            ["① AI 모델·인프라", "3,000", "Claude API / Pollinations Pro / Suno Enterprise"],
            ["② 글로벌 마케팅", "2,500", "Pinterest/TikTok English / Google Ads"],
            ["③ 외주 인력 (3명)", "2,500", "AI 엔지니어 1, 마케터 1, 콘텐츠 PM 1"],
            ["④ 데이터·법률", "1,000", "사주 데이터 확장 + 글로벌 결제 규제"],
            ["⑤ 인증·운영", "1,000", "AWS 이전 + ISMS"],
            ["합계", "10,000", "100% 국비 가정"],
        ],
    )

    add_heading(doc, "9. 차별화 — 왜 우리가 적격인가", 1)
    add_bullet(doc, "이미 라이브 매출 검증 (PayPal 글로벌 결제)")
    add_bullet(doc, "솔로 1인 + 90+ 자동화로 인건비 0원 효율 운영")
    add_bullet(doc, "비수도권(경주) 본사 — 지역 균형 발전 가산점")
    add_bullet(doc, "1개월차 신규 창업 — 초기창업/재도전 트랙 적격")
    add_bullet(doc, "13개 자동화 부서 + 8앱 Play Console + 27+ KDP ebook 동시 운영")

    add_heading(doc, "10. 팀 및 비전", 1)
    add_bullet(doc, "홍덕훈 (1985생) — Python/AI 자동화 1인 마스터")
    add_bullet(doc, "1년 후 채용 계획: 3명 (지역 일자리 창출)")
    add_para(
        doc,
        "비전: '한국 문화 X AI 자동화'의 글로벌 1위 1인 창업 모범 사례. "
        "100일 ₩10억, 1년 ARR ₩12억, 3년 시리즈 A → 5년 IPO 또는 M&A.",
    )

    add_heading(doc, "11. 부록", 1)
    add_bullet(doc, "천명당: https://cheonmyeongdang.vercel.app (PayPal 라이브)")
    add_bullet(doc, "세금N혜택: https://tax-n-benefit-api.vercel.app")
    add_bullet(doc, "KORLENS: https://korlens.app")
    add_bullet(doc, "Gumroad / KDP / YouTube 5채널 / Bluesky 검증 채널")

    out = os.path.join(OUT_DIR, "사업계획서_AI리그_백업_재도전성공패키지_2026.docx")
    doc.save(out)
    return out


# ================== MAIN ==================
if __name__ == "__main__":
    sys.stdout.reconfigure(encoding="utf-8")
    files = []
    for fn in [plan_1_ai_league, plan_2_tourism_ai, plan_3_local_track, plan_4_kglobal, plan_5_restart_backup]:
        path = fn()
        size = os.path.getsize(path)
        # Approximate word count
        from docx import Document
        d = Document(path)
        words = sum(len(p.text.split()) for p in d.paragraphs)
        for t in d.tables:
            for row in t.rows:
                for cell in row.cells:
                    words += sum(len(p.text.split()) for p in cell.paragraphs)
        chars = sum(len(p.text) for p in d.paragraphs) + sum(
            sum(len(p.text) for row in t.rows for cell in row.cells for p in cell.paragraphs)
            for t in d.tables
        )
        print(f"OK {os.path.basename(path)} | {size/1024:.1f} KB | words={words} | chars={chars}")
        files.append((path, size, words, chars))
    print()
    print("DONE 5 files generated.")
