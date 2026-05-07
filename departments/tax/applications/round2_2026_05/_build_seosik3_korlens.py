"""(서식3) 참가신청서 — 유형2 데이터 마케팅 컨설팅 (KORLENS) 작성완료 docx 생성.

기준 hwp: (서식3)참가신청서_(유형2)AI 기반 데이터 마케팅 컨설팅_(기업명).hwp
신청 사업: 2026 관광기업 데이터·AI 활용 지원 사업 (한국관광공사, 마감 5/20)
신청 회사: 쿤스튜디오 / KORLENS (한국 여행 AI 큐레이터)
"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r'C:\Users\hdh02\Desktop\cheonmyeongdang\departments\tax\applications\round2_2026_05\(작성완료)서식3_참가신청서_유형2_KORLENS_2026-05-07.docx'


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_kv(table, k, v, gray=True):
    row = table.add_row().cells
    row[0].text = k
    row[1].text = v
    if gray:
        set_cell_bg(row[0], 'F2F2F2')
    for c in row:
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in c.paragraphs:
            for r in p.runs:
                r.font.name = '맑은 고딕'
                r.font.size = Pt(10)


def kv_table(doc, kvs, col0=4.0, col1=12.0):
    t = doc.add_table(rows=0, cols=2)
    t.style = 'Table Grid'
    for k, v in kvs:
        add_kv(t, k, v)
    for row in t.rows:
        row.cells[0].width = Cm(col0)
        row.cells[1].width = Cm(col1)
    return t


def add_section(doc, num, title):
    h = doc.add_heading(f'{num}. {title}', level=1)
    return h


def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)

    # 제목
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('서식3').bold = True

    h = doc.add_heading('AI 기반 데이터 마케팅 컨설팅(그로스 해킹) 유형 참가신청서', level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph('※ 작성분량에는 제한이 없으나, 반드시 양식에 맞도록 작성하여 주시기 바랍니다.')
    p.runs[0].italic = True

    # ========== 일반 ==========
    kv_table(doc, [
        ('기업명 (사업자등록증과 동일)', '쿤스튜디오 (KunStudio)'),
        ('대표자명', '홍덕훈'),
        ('사업자등록번호', '552-59-00848'),
        ('근로자 수', '1명 (대표자)'),
        ('본점 소재지', '(38204) 경상북도 경주시 외동읍 제내못안길 25-52'),
        ('사업 참여 인력', '1명 (대표 홍덕훈 — Full-stack + 기획/마케팅/데이터)'),
        ('근무 소재지', '(38204) 경상북도 경주시 외동읍 제내못안길 25-52  *재택형 1인 기업'),
    ])

    # ========== 1. 사업현황 ==========
    add_section(doc, '1', '사업현황')
    doc.add_paragraph('가. 일반현황').runs[0].bold = True

    kv_table(doc, [
        ('서비스명', 'KORLENS — 한국 여행 AI 큐레이터'),
        ('관광기업 여부', '☑ 관광 분야 사업을 진행중   ☐ 관광 분야 사업을 진행하지 않음'),
        ('관광산업 연관성',
         'KORLENS는 외국인 한국 여행객을 대상으로 한국관광공사 TourAPI(국문관광정보), '
         '한국문화정보원(K-Data), 공공데이터포털 등의 관광 데이터를 AI(LLM)로 큐레이션해 '
         '맞춤형 여행 코스/맛집/숙박/액티비티를 추천하는 SaaS입니다. '
         '관광공사 콘텐츠 ID(contentId, contentTypeId)를 핵심 백엔드로 사용합니다.'),
        ('한국관광공사 관광기업 지원사업 수혜이력',
         '없음 (2026-04-01 신규 개업, 본 사업이 첫 신청)'),
        ('관광기업 혁신바우처 수혜이력', '없음'),
        ('사업구조',
         '☑ B2C (모바일 웹/앱: korlens.app)\n'
         '☑ B2B2C (호텔/여행사 화이트라벨 — 진행 중)\n'
         '☐ B2B   ☐ 기타'),
        ('2024년 주요성과 (매출액)', '해당없음 (2026-04-01 개업)'),
        ('2024년 주요성과 (월 평균 사용자)', '해당없음'),
        ('2025년 주요성과 (매출액)', '해당없음 (2026-04-01 개업)'),
        ('2025년 주요성과 (월 평균 사용자)', '해당없음'),
        ('직원 수 (2026-05 현재 기준)', '정규직 1명 (대표) / 계약직 0명'),
        ('투자유치 현황', '0백만원 / 누적 0백만원 (자본금 기반 자기자본 운영)'),
        ('서비스 운영 단계',
         '☑ 출시 3개월 이내 (2026-04 베타 런칭, 2026-05 글로벌 4개국어 정식 오픈)'),
        ('예정 상황',
         '☑ 운영 중\n'
         '☑ 신규 서비스 런칭 예정 (예상시기: 2026-06 호텔 화이트라벨 B2B 모듈)'),
        ('사용중인 사내 메신저', '☑ 카카오톡 (무료)   ☐ 네이버웍스   ☐ 슬랙'),
        ('서비스 대상 상세',
         '주요 사용자: ① 한국 방문 외국인 자유여행객(20~40대, 영어/일본어/중국어/한국어), '
         '② 한국 방문 첫 여행자(가이드북 대체 수요), '
         '③ K-콘텐츠 팬덤 기반 테마 여행객(K-pop/K-drama 성지순례).'),
        ('서비스 대상 구분',
         '☑ 외국인 (일본, 중국, 미국, 유럽 영어권, 동남아 — 4개국어 지원)\n'
         '☑ 내국인 (역방향: 한국인의 국내 여행 큐레이션)'),
        ('서비스 지역', '☑ 국내 (전국 17개 시·도)\n☑ 국외 (해외 사용자 접속 / 마케팅)'),
        ('운영 언어', '☑ 한국어   ☑ 영어   ☑ 일어   ☑ 중국어 (간체)'),
        ('서비스 소개',
         'KORLENS는 한국 여행객의 자연어 질문(예: "강남 야경 + 24시간 식당 코스 짜줘")을 '
         'LLM으로 해석해 한국관광공사 TourAPI 데이터를 실시간으로 큐레이션, '
         '맞춤 여행 코스를 PDF/카카오맵/구글맵 연동으로 제공하는 모바일 웹 SaaS입니다. '
         '핵심 기능: ① AI 코스 생성, ② 실시간 운영시간/혼잡도 체크, ③ 다국어 음성 길안내, '
         '④ 결제 연동(예약/티켓), ⑤ K-콘텐츠 성지 큐레이션.'),
        ('차별성/경쟁력',
         '① 4개국어 동시 LLM 응답 (경쟁사 Trazy/Klook은 검색 기반·LLM 미적용)\n'
         '② 한국관광공사 공식 데이터 직결 (정확성)\n'
         '③ 천명당(AI 사주 SaaS) 글로벌 4개국어 운영 노하우 이전 (이미 검증)\n'
         '④ K-콘텐츠 성지 데이터셋 자체 구축 (K-pop/K-drama 매핑 800+)\n'
         '⑤ 1인 기업 → 광고비/AI 솔루션 지원 시 즉시 투입 가능 (의사결정 0 lag)'),
        ('사용자의 서비스 이용 시나리오',
         '① 외국인이 인천공항 도착 → korlens.app 접속 (Google 검색 또는 광고 유입)\n'
         '② 자국어 선택 → "3박 4일 서울+부산 코스, 매운음식 좋아함" 입력\n'
         '③ 30초 내 LLM이 일자별 코스 + 지도 + 예약 링크 제공\n'
         '④ 코스 카카오맵/구글맵으로 내보내기 + PDF 다운로드\n'
         '⑤ 현지 이동 중 실시간 영업시간/혼잡도 푸시 알림'),
        ('핵심 목표 (3개 이내)',
         '① 외국인 회원가입 완료 (이메일/SNS OAuth)\n'
         '② AI 코스 생성 완료 + PDF 다운로드/공유\n'
         '③ 제휴 예약 클릭/전환 (Klook/Trip.com affiliate)'),
        ('기타 참고사항 및 자료',
         '대표 화면 캡처 / 데모 영상 별첨 가능. korlens.app 접속 시 즉시 체험 가능.'),
    ])

    # ========== 2. 개발환경 ==========
    add_section(doc, '2', '개발 환경 / 데이터 활용 현황')
    kv_table(doc, [
        ('현재 운영 중인 서비스 분류 (웹)', '☑ 웹서비스 있음 (https://korlens.app)'),
        ('서비스 개발 진행 방식 (웹)', '☑ 자체 제작 (대표 1인 풀스택)'),
        ('현재 운영 중인 서비스 분류 (앱)',
         '☑ 앱서비스 있음 (정확한 앱 명칭: KORLENS Mobile Web — PWA)'),
        ('앱서비스 분류',
         '☑ 하이브리드 (PWA Progressive Web App)\n'
         '※ 향후 React Native 네이티브 앱 빌드 예정 (2026 Q3)'),
        ('사용 중인 앱 개발 언어', '기타: TypeScript, Next.js 14, React 18'),
        ('사용 중인 앱 프레임워크', '☑ 하이브리드 (Next.js + Tailwind + Vercel Edge)'),
        ('현재 AI 도구 사용 여부', '☑ 도구가 있다'),
        ('사용 중인 AI 도구명',
         'OpenAI GPT-4o (LLM 코스 생성), Anthropic Claude 3.5 Sonnet (멀티모달 큐레이션), '
         'Google Vertex AI Gemini (다국어), Whisper (음성 인식), DeepL API (번역 보조)'),
        ('AI 도구 구독 현황',
         '☑ 유료 사용 — OpenAI Tier 4 / Anthropic Build Plan / Google Cloud / DeepL Pro'),
        ('AI 도구 활용 정도', '☑ 상시 활용 (실시간 LLM 응답이 핵심 서비스)'),
        ('AI 도구 활용 사례',
         '① 코스 생성 (LLM): 사용자 질문 → TourAPI 결과 → LLM 큐레이션\n'
         '② 다국어 응답: 4개국어 동시 LLM 출력\n'
         '③ 마케팅 소재 제작: GPT/Claude로 SNS 카피 자동 생성\n'
         '④ 데이터 분석: GA4 + Vercel Analytics 결과를 LLM이 일일 요약'),
        ('현재 분석 도구 사용', '☑ 분석 도구가 있다 (Google Analytics 4, Vercel Analytics, Microsoft Clarity)'),
        ('분석 도구 활용 정도', '상시 활용 (대시보드 일일 모니터링)'),
    ])

    # ========== 3. 디지털 마케팅 현황 ==========
    add_section(doc, '3', '디지털 마케팅 현황')
    kv_table(doc, [
        ('마케팅팀 업무소개',
         '대표 1인이 마케팅 전담 (개발 + 마케팅 동시). 현재 가장 큰 고민: '
         '① 외국인 타겟 광고 ROAS 측정 인프라 부재, '
         '② 4개국어 번역 콘텐츠의 채널별 성과 분리 분석 불가, '
         '③ 광고비 예산 한계로 데이터 마케팅 컨설팅 절실.'),
        ('2025년 마케팅 운영 성과',
         '쿤스튜디오는 2026-04-01 신규 개업으로 2025년 운영 성과 없음. '
         '2026-04~05 운영 채널: Pinterest, TikTok English, X(Twitter), Reddit Travel, '
         'Naver 블로그, Threads, Bluesky. 월 평균 광고비 0원 (모두 무료/affiliate 채널).'),
        ('2026년 비즈니스 목표 및 마케팅 전략',
         '① MAU: 외국인 5,000명 (영어 50%, 일본어 25%, 중국어 15%, 한국어 10%)\n'
         '② 매출: ₩30M (Klook/Trip.com affiliate + 호텔 화이트라벨 B2B)\n'
         '③ 마케팅 운영 방향: 광고비 1,000만원 = (a) Google Ads 외국인 타겟 60%, '
         '(b) Meta(IG/FB) 일본·동남아 25%, (c) Pinterest 영어권 15%. '
         '데이터 분석은 GA4 + 한국관광공사 데이터 결합으로 외국인 행동 패턴 분석.'),
        ('주요 확인 목표 지표',
         '① 외국인 신규 가입자 수 (월별)\n'
         '② AI 코스 생성 완료율 (%)\n'
         '③ Klook/Trip.com affiliate 클릭/전환율\n'
         '④ 4개국어별 체류 시간 / 이탈률\n'
         '⑤ 호텔 화이트라벨 데모 신청 수'),
        ('지표 개선 전략 (기존)',
         '4월 베타 운영 결과: SNS 무료 채널은 한국어 80%, 외국어 20% 유입 → '
         '외국인 타겟 도달이 핵심 과제. Pinterest 핀 영어/일본어 일 3핀 자동 발행, '
         'Reddit r/koreatravel 주 1회 가치 콘텐츠 게시, '
         'Klook AID(commission link) 인스타그램 설명 추가 — 5/6 시작.'),
    ])

    # ========== 4. 지원동기 및 참가목표 ==========
    add_section(doc, '4', '지원동기 및 참가목표')
    kv_table(doc, [
        ('데이터 마케팅 운영 현황 및 당면 과제',
         '현재 광고비 0원으로 무료 채널만 운영 중이라 외국인 타겟 도달 데이터를 충분히 '
         '수집하지 못하고 있습니다. 데이터 활용 희망 부분: '
         '① Google Ads + Meta Ads 외국인 타겟 광고 성과 지표(ROAS, CPA) 확인 및 최적화, '
         '② GA4와 한국관광공사 TourAPI 호출 로그 결합 분석을 통한 사용자 코스 선호도 파악, '
         '③ 4개국어별 사용자 여정 차이 분석 → 언어별 LLM 프롬프트 최적화.'),
        ('기대 지원사항',
         '본 사업 광고비 1,000만원 + AI 솔루션 100만원 활용 계획:\n'
         '(1) Google Ads 600만원 — 외국인 검색어("Korea travel", "Seoul itinerary") 타겟\n'
         '(2) Meta Ads 250만원 — 일본/동남아 IG 인플루언서 페어링\n'
         '(3) Pinterest Ads 150만원 — 영어권 시각형 광고\n'
         '(4) AI 솔루션 100만원 — Mixpanel/Amplitude 6개월 구독 + GA4 BigQuery export 셋업\n'
         '컨설팅 영역: ROAS 측정 인프라 구축, 외국인 행동 funnel 분석, 광고 크리에이티브 A/B 테스트.'),
        ('AI 활용 현황',
         '현재 OpenAI GPT-4o + Claude 3.5 Sonnet + Gemini 멀티 LLM 라우팅으로 LLM 응답 생성. '
         '어려움: ① 광고 크리에이티브 자동 생성 도구 선택 조언 필요, '
         '② Mixpanel 같은 product analytics 도입 결정 필요 (GA4 한계).'),
        ('AI 활용 (희망)',
         '본 지원사업 통해 진행 희망: ① AI 광고 크리에이티브 자동화(AdCreative.ai 등) 도입 컨설팅, '
         '② Mixpanel + AI 행동 분석 셋업, ③ 4개국어 LLM 프롬프트 최적화 워크숍.'),
        ('수행을 위한 내부 인력 지원 계획',
         '대표 홍덕훈 1인이 모든 컨설팅 회의 직접 참여. '
         '주 5일 9~22시 화상 미팅 가능 (시간 제약 없음). '
         '컨설턴트 요청 자료/데이터 export 24시간 이내 제공 가능 (1인 의사결정 0 lag).'),
        ('기타 지원 계획 (인력 외)',
         '① 모든 분석 데이터 (GA4, Vercel Analytics, TourAPI 호출 로그) 무제한 접근 권한 제공\n'
         '② A/B 테스트 즉시 코드 반영 (대표가 직접 개발)\n'
         '③ 컨설팅 결과를 KORLENS 공식 블로그/유튜브로 공개해 사례 확산'),
    ])

    # ========== 5. 마케팅 인력 ==========
    add_section(doc, '5', '마케팅 및 서비스 운영 인력 구성 현황')
    p = doc.add_paragraph(
        '마케팅, 그로스 해킹 관련 인력 보유 현황:\n'
        '- 대표 홍덕훈 1명 (마케팅 + 그로스 + 콘텐츠 + 퍼포먼스 통합 담당)\n'
        '※ 1인 기업으로 모든 마케팅 역할을 대표가 직접 수행. '
        '본 컨설팅을 통해 1인 기업도 데이터 기반 그로스가 가능함을 입증할 계획.'
    )

    # ========== 6. 데이터 인력 ==========
    add_section(doc, '6', '데이터 및 개발 관련 참여인력 현황')
    p = doc.add_paragraph(
        '데이터 수집 작업 가능한 인력 보유 현황:\n'
        '- 대표 홍덕훈 1명 (Full-stack: 기획 + Frontend + Backend + Data + AI)\n'
        '  ・ Frontend: Next.js 14, React 18, TypeScript, Tailwind\n'
        '  ・ Backend: Node.js, Vercel Edge Functions, Supabase, PostgreSQL\n'
        '  ・ Data: GA4, BigQuery, Mixpanel(예정), TourAPI 직접 통합\n'
        '  ・ AI: OpenAI/Anthropic/Google LLM API 라우팅 시스템 자체 구축\n'
        '※ 1인 풀스택으로 데이터 수집/분석/개발 모두 자체 수행 가능.'
    )

    # ========== 7. 사업 책임자 ==========
    add_section(doc, '7', '사업 책임자 정보')
    kv_table(doc, [
        ('성명', '홍덕훈 (Hong Deokhoon)'),
        ('직위/직책', '대표'),
        ('연락처', '010-4244-6992'),
        ('이메일', 'ghdejr11@gmail.com'),
        ('참여 기업 재직기간', '1개월 (2026-04-01 개업~)'),
        ('총 경력', '소프트웨어 개발 / 기획 5년 이상'),
        ('직무', '대표 (총괄 — 기획·개발·마케팅·데이터)'),
        ('담당업무 소개',
         '쿤스튜디오 대표로서 천명당(AI 사주 SaaS, 4개국어), 세금N혜택(소상공인 세무 자동화), '
         'KORLENS(한국 여행 AI 큐레이터) 3개 라인업 전체를 기획/개발/운영. '
         '본 컨설팅의 책임자이자 실무자로 모든 회의 직접 참여.'),
        ('주요 이력',
         '- 천명당(cheonmyeongdang.com) 개발·런칭 (4개국어 SaaS, 2026-05 글로벌 정식 오픈)\n'
         '- KORLENS(korlens.app) 개발·런칭 (한국 여행 AI 큐레이터)\n'
         '- 세금N혜택 개발 (소상공인 세무 자동화 SaaS)\n'
         '- LLM 기반 멀티 AI 라우팅 시스템 자체 구축 (OpenAI + Anthropic + Google)\n'
         '- Google Analytics 4 + Vercel Analytics 기반 사용자 행동 분석'),
    ])

    # ========== 8. 담당 참여 인력 ==========
    add_section(doc, '8', '담당 참여 인력 정보')
    kv_table(doc, [
        ('성명', '홍덕훈 (대표 = 책임자 = 실무자, 1인 기업)'),
        ('직위/직책', '대표'),
        ('연락처', '010-4244-6992'),
        ('이메일', 'ghdejr11@gmail.com'),
        ('재직기간', '1개월 (2026-04-01 개업~)'),
        ('직무', '데이터 마케팅 + 그로스 해킹 + 개발 통합'),
        ('담당업무 소개',
         '본 컨설팅의 모든 회의·실행·보고를 단독 담당. '
         '컨설턴트와 직접 소통 + 즉시 코드 반영(1인 풀스택).'),
        ('주요 이력', '상기 사업 책임자 정보와 동일'),
    ])

    # ========== 서명 ==========
    doc.add_paragraph()
    p = doc.add_paragraph(
        '상기 본인은 「AI 기반 데이터 마케팅 컨설팅(그로스 해킹) 유형」에 참여하고자 '
        '위와 같이 신청서를 제출하며, 위 기재사항이 사실과 틀림없음을 확인합니다.'
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('\n2026년 05월 07일\n\n대표자: 홍 덕 훈   (인)')
    r.bold = True
    r.font.size = Pt(11)

    doc.save(OUT)
    print(f'[OK] {OUT}')
    print(f'  size: {os.path.getsize(OUT):,} bytes')


if __name__ == '__main__':
    main()
