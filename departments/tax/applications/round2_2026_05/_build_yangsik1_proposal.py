"""(양식1) 2026 관광데이터 활용 공모전 — 앱 개발 부문 제안서 작성완료 docx.

- 원본 hwp: (양식1)『2026 관광데이터 활용 공모전』 제안서_팀명(대표자명).hwp
- 분량 제한: 5페이지 이내, 12pt 이상
- 팀명: 쿤스튜디오 / 대표자명: 홍덕훈
- 제출 파일명: 『2026 관광데이터 활용 공모전』 제안서_쿤스튜디오(홍덕훈).docx
"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r'D:\cheonmyeongdang\departments\tax\applications\round2_2026_05\(작성완료)양식1_관광데이터공모전_제안서_쿤스튜디오_홍덕훈_2026-05-07.docx'


def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(12)  # 12pt 이상 필수

    # 헤더
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('[참고] 제안서 양식').bold = True

    h = doc.add_heading('『2026 관광데이터 활용 공모전』 앱 개발 부문 제안서', level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('팀명: 쿤스튜디오  /  대표자명: 홍덕훈')
    r.bold = True
    r.font.size = Pt(13)

    doc.add_paragraph(
        '※ 작성항목 및 양식 수정 금지, 모든 항목 작성 필수, 이미지(사진)·표 등 삽입 가능. '
        '분량: 5페이지 이내(이미지·표 등 일체 포함), 글자 포인트는 12포인트 이상.'
    )

    # ========== 1) 서비스 기획배경 및 필요성 ==========
    doc.add_heading('1) 서비스 기획배경 및 필요성', level=1)

    p = doc.add_paragraph()
    p.add_run('● 서비스 기획 배경\n').bold = True
    doc.add_paragraph(
        '한국은 K-pop·K-drama·K-food 글로벌 인기 확산으로 외국인 방문객이 '
        '2024년 1,600만 명을 돌파했으나, 외국인 자유여행객의 80% 이상이 '
        '"서울·강남·명동" 등 수도권 핫스팟에만 집중되고, 인구 감소지역(전국 89개 시·군) '
        '관광 정보를 영어/일본어/중국어로 접근할 수 있는 통합 서비스가 부재한 상황이다. '
        '한편 한국관광공사 TourAPI는 양질의 데이터를 무료 제공하지만 다국어 큐레이션·'
        '맞춤형 추천 인터페이스가 없어 외국인 사용자가 직접 활용하기 어렵다. '
        '쿤스튜디오는 천명당(AI 사주 SaaS, 4개국어 글로벌 운영) 개발 경험을 바탕으로 '
        'TourAPI 데이터를 LLM(GPT-4o + Claude 3.5)으로 큐레이션하는 '
        '"KORLENS — 한국 여행 AI 큐레이터"를 기획했다.'
    )

    p = doc.add_paragraph()
    p.add_run('● 서비스 필요성\n').bold = True
    doc.add_paragraph(
        '① 외국인 자유여행객의 인구 감소지역 분산 유입을 통한 관광 균형발전 기여 — '
        '경주·안동·통영·여수 등 인구 감소지역의 우수한 관광 자원이 외국어 정보 부족으로 '
        '외면받는 현실 해결.\n'
        '② 4개국어(한국어·영어·일본어·중국어) LLM 응답으로 외국인이 자국어로 자연스럽게 '
        '여행 코스를 설계 가능 — 기존 Visit Korea 앱은 정적 정보 검색만 제공.\n'
        '③ 한국관광공사 TourAPI(국문관광정보), 한국문화정보원, 공공데이터포털을 단일 '
        '인터페이스로 통합해 외국인이 1개 앱으로 모든 한국 여행 정보 접근 가능.\n'
        '④ 쿤스튜디오는 천명당(4개국어 사주 SaaS) 운영 경험으로 다국어 LLM 응답 노하우와 '
        '인프라(Vercel Edge + Supabase + 멀티 LLM 라우팅)를 보유 → 즉시 구현 가능.'
    )

    # ========== 2) 서비스 개요 ==========
    doc.add_heading('2) 서비스 개요', level=1)

    p = doc.add_paragraph()
    p.add_run('● 기획 서비스 소개 (한 줄): ').bold = True
    p.add_run('LLM 기반 4개국어 한국 여행 AI 큐레이터 — 인구 감소지역 우선 추천 알고리즘 탑재')

    p = doc.add_paragraph()
    p.add_run('● 기획 서비스 주요 기능\n').bold = True
    doc.add_paragraph(
        '① AI 코스 생성: 사용자 자연어 질문(예: "Korea travel 5 days, K-drama spots") → '
        'TourAPI 호출 → LLM 큐레이션 → 일자별 코스 + 지도 + 예약 링크 (30초 내)\n'
        '② 인구 감소지역 우선 추천: 경주·안동·통영·여수·강릉 등 89개 인구 감소지역 데이터셋을 '
        '서울·부산 코스에 자동 끼워넣어 분산 유도 (관광 균형발전)\n'
        '③ 4개국어 동시 응답: 한국어·영어·일본어·중국어 — 사용자 선택 즉시 자국어 LLM 응답\n'
        '④ K-콘텐츠 성지 매핑: K-pop MV 촬영지, K-drama 명장면 장소 800+ 데이터셋 자체 구축\n'
        '⑤ 카카오맵·구글맵 연동 코스 내보내기 + PDF 다운로드 + 음성 길안내 (Whisper)\n'
        '⑥ Klook·Trip.com Affiliate 자동 링크 → 실시간 티켓·숙박 예약 전환'
    )

    p = doc.add_paragraph()
    p.add_run('● 서비스 차별성\n').bold = True
    doc.add_paragraph(
        '① 4개국어 LLM 동시 응답 — 경쟁 서비스(Trazy, Klook, KKday)는 검색 기반·LLM 미적용\n'
        '② 한국관광공사 TourAPI 직결 (정확성·실시간성) — 경쟁사는 자체 DB로 정보 노후화\n'
        '③ 인구 감소지역 우선 추천 알고리즘 (관광 균형발전 정책 부합)\n'
        '④ K-콘텐츠 성지 자체 데이터셋 800+ — 글로벌 K-팬덤 타겟 차별화\n'
        '⑤ 1인 풀스택 기업 — Vercel Edge + 멀티 LLM 라우팅으로 운영비 월 ₩300K 이하 (저비용 지속가능)'
    )

    # ========== 3) 데이터 활용 방식 ==========
    doc.add_heading('3) 한국관광공사 데이터 활용 방식', level=1)

    p = doc.add_paragraph()
    p.add_run('● 활용 예정 한국관광공사 데이터 (필수)\n').bold = True
    doc.add_paragraph(
        '① 한국관광공사 국문관광정보 서비스 (TourAPI 4.0)\n'
        '   - areaBasedList (지역기반 관광지·맛집·숙박)\n'
        '   - searchKeyword (키워드 검색)\n'
        '   - detailIntro (상세소개 — 운영시간, 휴무일, 입장료)\n'
        '   - detailImage (대표 이미지)\n'
        '② 한국관광공사 영문/일문/중문 관광정보 서비스 (다국어 TourAPI)\n'
        '③ 한국관광공사 인구 감소지역 관광지 데이터셋 (89개 시·군)\n'
        '④ 한국관광공사 코스 정보 서비스 (제안 코스 25유형)'
    )

    p = doc.add_paragraph()
    p.add_run('● 데이터 활용 방식 (기술적 구현)\n').bold = True
    doc.add_paragraph(
        '① 사용자 자연어 질문 → LLM(GPT-4o function calling)이 TourAPI 호출 파라미터 '
        '자동 생성 (areaCode, contentTypeId, keyword 등)\n'
        '② TourAPI 응답(JSON) → LLM이 사용자 선호도/제약(예산, 일자, 동반자) 반영해 큐레이션\n'
        '③ 인구 감소지역 가중치 알고리즘: 코스 내 1개 이상 인구 감소지역 자동 포함 시도\n'
        '④ 다국어 TourAPI 동시 호출 → 사용자 언어로 결과 표시\n'
        '⑤ Vercel Edge 캐싱 (1시간 TTL)으로 TourAPI 호출 비용 최소화 + 응답 속도 향상\n'
        '⑥ Supabase에 호출 로그 저장 → 인기 코스/지역 분석 → 한국관광공사에 인사이트 환원'
    )

    # ========== 4) 서비스 발전 방향 ==========
    doc.add_heading('4) 서비스 발전 방향', level=1)

    p = doc.add_paragraph()
    p.add_run('● 개발 서비스 향후 발전방향 (1년 로드맵)\n').bold = True
    doc.add_paragraph(
        'Q3 2026: 호텔/여행사 화이트라벨 B2B 모듈 (호텔 객실에서 QR로 KORLENS 접속)\n'
        'Q4 2026: React Native 네이티브 앱 출시 (iOS·Android 동시) + 오프라인 모드\n'
        'Q1 2027: AR 길안내 (카메라로 비추면 LLM이 음성으로 설명)\n'
        'Q2 2027: 태국어·베트남어·인도네시아어 추가 (동남아 여행객 증가 대응)\n'
        'Q3 2027: 인구 감소지역 지자체 직접 제휴 (경주시·안동시·여수시 공식 파트너)'
    )

    p = doc.add_paragraph()
    p.add_run('● 기대효과\n').bold = True
    doc.add_paragraph(
        '【정량】\n'
        '- 1년 내 외국인 MAU 50,000명 (영어 50%, 일본어 25%, 중국어 15%, 한국어 10%)\n'
        '- 인구 감소지역 송출 트래픽 30% 이상 (코스 내 1개 이상 포함률 기준)\n'
        '- 매출 ₩300M (Affiliate + 화이트라벨 B2B + 지자체 제휴)\n'
        '- TourAPI 호출 100만 회/월 → 한국관광공사 데이터 ROI 입증 사례\n\n'
        '【정성】\n'
        '- 외국인 자유여행객의 한국 여행 진입 장벽 해소 → 관광 강국 위상 강화\n'
        '- 인구 감소지역 균형 송출 → 지방소멸 대응 정책 부합\n'
        '- 한국관광공사 공공데이터의 민간 활용 모범 사례 → 다른 데이터 활용 촉진\n'
        '- 1인 기업이 공공데이터 + AI로 글로벌 SaaS 구축한 성공 케이스 → 청년 창업 자극'
    )

    # 마지막
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('\n2026년 05월 07일\n\n팀명: 쿤스튜디오  /  대표자: 홍 덕 훈')
    r.bold = True
    r.font.size = Pt(13)

    doc.save(OUT)
    print(f'[OK] {OUT}')
    print(f'  size: {os.path.getsize(OUT):,} bytes')


if __name__ == '__main__':
    main()
