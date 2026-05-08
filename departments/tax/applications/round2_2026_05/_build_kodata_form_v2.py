"""KoDATA 양식 v2 — 5/7 통화 핵심 반영.

5/7 14:32 통화 핵심:
  1. 모든 서명란 서명 (3곳: 기업명 옆 / 기업개요표 / 정보제공 동의서)
  2. 대표자 주민번호 뒷자리 기재 → docx에 placeholder 명시
  3. 사업자등록증 PDF 깨짐 → JPG 변환 첨부 (별도 task)
  4. 별도 docx 작성 OK (KoDATA 안내)

변경:
  - 기존 (작성완료) docx → v2로 복제
  - 표 [4] r3 (대표자 주민등록번호) 마스킹 → 사용자 직접 입력 placeholder
  - 표 [4] r4 (서명/날인) → 사용자 서명 이미지 삽입 안내 placeholder
  - 표 [0] r0 (기업명 옆) 서명 placeholder 추가
  - 표 [2] r0 (기업개요표 기업명) 서명 placeholder 추가
  - "[작성: 2026-05-07 / 쿤스튜디오 대표 홍덕훈]" 마지막 paragraph 추가
"""
import os
import sys
import shutil
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import RGBColor, Pt

ROUND2 = r'D:\cheonmyeongdang\departments\tax\applications\round2_2026_05'
SRC = os.path.join(ROUND2, '(작성완료)한국관광공사_기업정보_등록의뢰서_쿤스튜디오_2026-05-07.docx')
DST = os.path.join(ROUND2, '(작성완료_v2)한국관광공사_기업정보_등록의뢰서_쿤스튜디오_2026-05-07.docx')

# 사용자 직접 입력 placeholder (눈에 띄도록 [[]] + 빨강)
RED = RGBColor(0xC0, 0x00, 0x00)


def set_cell_text(cell, text, *, bold=False, red=False):
    """셀 내용 전체 교체 (단일 paragraph)."""
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    if bold:
        run.bold = True
    if red:
        run.font.color.rgb = RED
        run.bold = True


def append_run(paragraph, text, *, red=False, bold=False, size=None):
    run = paragraph.add_run(text)
    if red:
        run.font.color.rgb = RED
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    return run


def main():
    print('=== KoDATA docx v2 빌드 ===')
    print(f'  SRC: {SRC}')
    print(f'  DST: {DST}')
    if not os.path.isfile(SRC):
        raise FileNotFoundError(SRC)

    shutil.copy2(SRC, DST)
    doc = Document(DST)

    tables = doc.tables
    print(f'  Tables: {len(tables)}')
    # table 0 = 입력 대상기업 (기업명/사업자번호)
    # table 1 = 담당자
    # table 2 = 기업개요표
    # table 3 = 동의서 체크박스
    # table 4 = 서명 (기업체명/사업자번호/대표자/주민번호/서명)

    # ──────────────────────────────────────────
    # ① 표[0] (입력 대상기업) — 기업명 옆 서명란
    # ──────────────────────────────────────────
    # 기존 r0 = ['기업명', '쿤스튜디오 (KunStudio)']
    # 셀 r0c1 끝에 서명 placeholder 추가
    t0_r0c1 = tables[0].rows[0].cells[1]
    p = t0_r0c1.add_paragraph()
    append_run(p, '서명: ', bold=True)
    append_run(p, '[[XXX 자필서명 또는 이미지 삽입 ①]]', red=True, bold=True)
    print('  [OK] 표0 r0c1: 기업명 옆 서명 placeholder 추가')

    # ──────────────────────────────────────────
    # ② 표[2] (기업 개요표) — 기업명 옆 서명란
    # ──────────────────────────────────────────
    t2_r0c1 = tables[2].rows[0].cells[1]
    p = t2_r0c1.add_paragraph()
    append_run(p, '서명: ', bold=True)
    append_run(p, '[[XXX 자필서명 또는 이미지 삽입 ②]]', red=True, bold=True)
    print('  [OK] 표2 r0c1: 기업개요표 기업명 옆 서명 placeholder 추가')

    # ──────────────────────────────────────────
    # ③ 표[4] r3 (대표자 주민등록번호) — 사용자 직접 입력
    # ──────────────────────────────────────────
    # 기존: '850813-1******* (마스킹, 별도 본인확인 시 제공)'
    # 변경: 앞 6자리는 그대로, 뒷 7자리는 placeholder 명시
    t4 = tables[4]
    set_cell_text(t4.rows[3].cells[1],
                  '850813-[[XXX 사용자 직접 입력 — 주민번호 뒷 7자리]]',
                  bold=True, red=True)
    print('  [OK] 표4 r3: 주민번호 뒷자리 placeholder')

    # ──────────────────────────────────────────
    # ④ 표[4] r4 (서명/날인) — 정보제공 동의서 서명
    # ──────────────────────────────────────────
    # 기존: '(전자서명) 홍덕훈 / 2026-05-07'
    # 변경: 자필서명 또는 이미지 삽입 placeholder
    set_cell_text(t4.rows[4].cells[1],
                  '[[XXX 자필서명 또는 이미지 삽입 ③ — 정보제공 동의서]]  /  2026-05-07',
                  bold=True, red=True)
    print('  [OK] 표4 r4: 정보제공 동의서 서명 placeholder')

    # ──────────────────────────────────────────
    # ⑤ 마지막에 작성자 명시 단락 추가
    # ──────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph()
    append_run(p, '─' * 30, bold=True)
    p = doc.add_paragraph()
    append_run(p, '[v2 정정 — 2026-05-07 14:32 KoDATA 통화 안내 반영]', red=True, bold=True, size=11)
    p = doc.add_paragraph()
    append_run(p, '  ① 모든 서명란 3곳: 사용자 자필서명 또는 이미지 삽입 필요 (위 [[XXX]] 위치)', size=10)
    p = doc.add_paragraph()
    append_run(p, '  ② 대표자 주민번호 뒷자리: 사용자 직접 입력 (위 [[XXX]] 위치, 보안상 자동 X)', size=10)
    p = doc.add_paragraph()
    append_run(p, '  ③ 사업자등록증: 별도 첨부 (사업자등록증_쿤스튜디오_2026-05-07.jpg)', size=10)
    p = doc.add_paragraph()
    append_run(p, '작성: 2026-05-07 / 쿤스튜디오 대표 홍덕훈', bold=True, size=11)
    print('  [OK] 마지막 단락: 작성자 명시 + v2 안내')

    doc.save(DST)
    print()
    print(f'  완료: {DST}')
    print(f'  크기: {os.path.getsize(DST):,} bytes')


if __name__ == '__main__':
    main()
