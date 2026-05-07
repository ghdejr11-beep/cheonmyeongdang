"""KoDATA 기업정보 등록의뢰서 docx fallback 생성 (2026-05-07).

원본 hwp 양식이 빈양식 회신을 받음 → 동일 정보를 docx로 채워서 재발송.
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r'C:\Users\hdh02\Desktop\cheonmyeongdang\departments\tax\applications\round2_2026_05\(작성완료)한국관광공사_기업정보_등록의뢰서_쿤스튜디오_2026-05-07.docx'

# 기업 정보
INFO = {
    '기업명': '쿤스튜디오 (KunStudio)',
    '사업자등록번호': '552-59-00848',
    '대표자': '홍덕훈',
    '대표자 생년월일': '1985-08-13',
    '서류 제출자(담당자)': '홍덕훈 (대표)',
    '담당자 전화번호': '010-4244-6992',
    '회사 대표전화번호': '070-8018-7832',
    '회사 대표팩스번호': '없음 (1인 기업)',
    '문자메시지수신요청': '예 (010-4244-6992로 발송 부탁드립니다)',
    '휴대전화번호(업무용)': '010-4244-6992',
    '이메일': 'ghdejr11@gmail.com',
    '법인등록번호': '해당없음 (개인사업자)',
    '본사주소': '(38204) 경상북도 경주시 외동읍 제내못안길 25-52',
    '사업장주소': '(38204) 경상북도 경주시 외동읍 제내못안길 25-52',
    '홈페이지': 'https://cheonmyeongdang.com / https://korlens.app',
    '업종(업태)': '정보통신업',
    '업종(종목)': '응용 소프트웨어 개발 및 공급업',
    '주요제품': '천명당(AI 사주 SaaS, 4개국어 글로벌 서비스), 세금N혜택(소상공인 세무 자동화), KORLENS(한국 여행 AI 큐레이터)',
    '설립일자(개업일자)': '2026-04-01 (자동입력)',
    '상시종업원수': '1명 (대표자, 자동입력)',
    '과세유형': '간이과세자',
    '사업 형태': '개인사업자',
}

ATTACHMENTS = [
    '기업개요표 (별도 양식 작성)',
    '사업자등록증명서 (개인 사업자)',
    '신용정보 제공 및 활용 동의서 (서명 포함)',
]

REASON = (
    '한국관광공사 「2026 관광기업 데이터·AI 활용 지원 사업」(KORLENS, 마감 2026.05.20 18:00) '
    '및 「2026 관광데이터 활용 공모전」 신청 자격 확보를 위한 한국관광산업포털(투어라즈) 가입.'
)


def add_table_row(table, label, value, bold_label=True):
    row = table.add_row().cells
    row[0].text = label
    row[1].text = value
    for c in row:
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for p in row[0].paragraphs:
        for r in p.runs:
            r.bold = bold_label
            r.font.name = '맑은 고딕'
    for p in row[1].paragraphs:
        for r in p.runs:
            r.font.name = '맑은 고딕'


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def main():
    doc = Document()
    # 폰트 기본
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)
    rPr = style.element.rPr
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # 제목
    h = doc.add_heading('한국관광공사 기업정보 등록의뢰서', level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run('작성일: 2026년 05월 07일')
    r.font.name = '맑은 고딕'
    r.font.size = Pt(10)

    # ========== 1. 입력 대상기업 ==========
    doc.add_heading('1. 입력 대상기업', level=1)
    t1 = doc.add_table(rows=0, cols=2)
    t1.style = 'Table Grid'
    t1.autofit = False
    fields_1 = [
        '기업명', '사업자등록번호',
    ]
    for k in fields_1:
        add_table_row(t1, k, INFO[k])

    # ========== 2. 서류 제출자 / 담당자 ==========
    doc.add_heading('2. 서류 제출자 / 담당자 정보', level=1)
    t2 = doc.add_table(rows=0, cols=2)
    t2.style = 'Table Grid'
    fields_2 = [
        '서류 제출자(담당자)',
        '담당자 전화번호',
        '회사 대표전화번호',
        '회사 대표팩스번호',
        '문자메시지수신요청',
        '휴대전화번호(업무용)',
        '이메일',
    ]
    for k in fields_2:
        add_table_row(t2, k, INFO[k])

    p = doc.add_paragraph()
    r = p.add_run('※ 진행현황을 휴대전화 문자메시지로 알려드리니, 반드시 기재합니다. '
                  '팩스발송 후 3시간 이후에도 문자메세지를 받지 못하면, 재발송해 주시기 바랍니다.')
    r.italic = True
    r.font.size = Pt(9)

    # ========== 3. 한국관광공사 기업 개요표 ==========
    doc.add_heading('3. 한국관광공사 기업 개요표', level=1)
    t3 = doc.add_table(rows=0, cols=2)
    t3.style = 'Table Grid'
    fields_3 = [
        '기업명',
        '법인등록번호',
        '대표자',
        '대표자 생년월일',
        '사업자등록번호',
        '홈페이지',
        '본사주소',
        '사업장주소',
        '업종(업태)',
        '업종(종목)',
        '주요제품',
        '설립일자(개업일자)',
        '상시종업원수',
        '과세유형',
        '사업 형태',
    ]
    for k in fields_3:
        add_table_row(t3, k, INFO[k])

    # ========== 4. 등록 사유 ==========
    doc.add_heading('4. 등록 사유', level=1)
    p = doc.add_paragraph(REASON)
    p.runs[0].font.name = '맑은 고딕'

    # ========== 5. 첨부서류 ==========
    doc.add_heading('5. 첨부서류 (필수 제출서류)', level=1)
    for a in ATTACHMENTS:
        doc.add_paragraph(f'☑  {a}', style='List Bullet')

    # ========== 6. 신용정보 동의 ==========
    doc.add_heading('6. 기업(신용)정보 및 개인(신용)정보 수집·활용·제공·조회 동의서', level=1)
    p = doc.add_paragraph()
    r = p.add_run('한국평가데이터(주) (이하 "당사"라 함) 귀중')
    r.bold = True

    doc.add_paragraph(
        '본인은 본 동의서의 내용을 이해하였으며, 기업(신용)정보 및 개인(신용)정보의 '
        '수집·활용·제공·조회 동의에 관한 내용에 대해 자세히 확인하였습니다.'
    )

    # 동의 항목
    consents = [
        ('1) 기업(신용)정보 수집·조회 동의 (필수)', '☑ 동의함    ☐ 동의하지 않음'),
        ('2) 개인(신용)정보 수집·활용 동의 (필수)', '☑ 동의함    ☐ 동의하지 않음'),
        ('3) 개인(신용)정보 제공 동의 (필수)', '☑ 동의함    ☐ 동의하지 않음'),
        ('4) 개인(신용)정보 조회 동의 (필수)', '☑ 동의함    ☐ 동의하지 않음'),
        ('5) 고유식별정보 수집·제공 동의 (필수, 주민등록번호 포함)', '☑ 동의함    ☐ 동의하지 않음'),
        ('6) 개인(신용)정보 수집·활용 동의 (선택)', '☑ 동의함    ☐ 동의하지 않음'),
    ]
    tc = doc.add_table(rows=0, cols=2)
    tc.style = 'Table Grid'
    for k, v in consents:
        add_table_row(tc, k, v)

    # ========== 7. 서명 ==========
    doc.add_heading('7. 서명 (Signature)', level=1)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('2026년 05월 07일')
    r.bold = True
    r.font.size = Pt(11)

    ts = doc.add_table(rows=0, cols=2)
    ts.style = 'Table Grid'
    add_table_row(ts, '기업체명', '쿤스튜디오 (KunStudio)')
    add_table_row(ts, '사업자등록번호', '552-59-00848')
    add_table_row(ts, '대표자 성명', '홍 덕 훈')
    add_table_row(ts, '대표자 주민등록번호', '850813-1******* (마스킹, 별도 본인확인 시 제공)')
    add_table_row(ts, '서명 / 날인', '(전자서명) 홍덕훈 / 2026-05-07')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('\n*  본 의뢰서는 한국관광공사 「2026 관광기업 데이터·AI 활용 지원 사업」 '
                  '신청을 위한 한국관광산업포털(투어라즈) 가입용으로 제출됩니다.')
    r.italic = True
    r.font.size = Pt(9)

    # 발송 정보
    doc.add_heading('8. 보내실 곳 / 문의처', level=1)
    p = doc.add_paragraph(
        '▶ 서류제출 메일: find@kodata.co.kr\n'
        '▶ 메일제목: 한국관광공사-쿤스튜디오\n'
        '▶ 진행 문의: 한국평가데이터 기업정보센터 02-3279-6500\n'
        '▶ 한국관광공사 공공사업부: 02-3215-2543'
    )

    # 컬럼 너비 조정
    for table in [t1, t2, t3, tc, ts]:
        for row in table.rows:
            row.cells[0].width = Cm(5.0)
            row.cells[1].width = Cm(11.0)

    doc.save(OUT)
    print(f'[OK] {OUT}')
    print(f'  size: {os.path.getsize(OUT):,} bytes')


if __name__ == '__main__':
    main()
