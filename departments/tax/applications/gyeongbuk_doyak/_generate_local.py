"""
2026 소상공인 도약 지원사업 — 로컬기업 육성 (지역 자원 기반)
사업계획서 작성: KORLENS 경주 특화 + 천명당 전통문화
"""
import sys
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

doc = Document()
style = doc.styles["Normal"]
style.font.name = "맑은 고딕"
style.font.size = Pt(10)


def h1(text):
    p = doc.add_heading(text, level=1)
    if p.runs: p.runs[0].font.name = "맑은 고딕"
    return p


def h2(text):
    p = doc.add_heading(text, level=2)
    if p.runs: p.runs[0].font.name = "맑은 고딕"
    return p


def para(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "맑은 고딕"; r.font.size = Pt(10)
    if bold: r.bold = True
    return p


def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    for r in p.runs:
        r.font.name = "맑은 고딕"; r.font.size = Pt(10)
    return p


# ─── 표지 ───
t = doc.add_heading("", 0)
r = t.add_run("2026년 소상공인 도약 지원 사업\n(로컬기업 육성 — 지역자원 기반)\n사업계획서")
r.font.name = "맑은 고딕"; r.font.size = Pt(18)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
doc.add_paragraph()

info = doc.add_table(rows=6, cols=2); info.style = "Light Grid"
rows = [
    ("상호", "쿤스튜디오 (경주 본사)"),
    ("대표자", "홍덕훈 (1985-02-03)"),
    ("사업자등록번호", "552-59-00848"),
    ("사업장", "경상북도 경주시 외동읍 제내못안길 26-62"),
    ("핵심 아이템", "KORLENS 경주 특화 AI 관광 큐레이션 + 천명당 K-전통문화 체험"),
    ("신청 유형", "로컬기업 육성 (지역성 50점)"),
]
for i, (k, v) in enumerate(rows):
    info.rows[i].cells[0].text = k
    info.rows[i].cells[1].text = v
doc.add_page_break()

# ─── 1. 기업 일반 현황 ───
h1("1. 기업 일반 현황")
t2 = doc.add_table(rows=8, cols=2); t2.style = "Light Grid"
rows = [
    ("업태·종목", "정보통신업 / 응용 소프트웨어 개발 및 공급업"),
    ("개업일자", "2026-04-01 (창업 20일차)"),
    ("본사 소재지", "경상북도 경주시 외동읍 (경주 관광 1시간권역)"),
    ("과세유형", "간이과세자 (소상공인 해당)"),
    ("상시 근로자", "1명 (대표자)"),
    ("지역 연결성", "경주 본사 + 경상북도 관광공사 공모전 제출(4/18)"),
    ("주요 서비스", "KORLENS(관광), 천명당(사주·전통문화), 세금N혜택, 보험다보여, KDP 전자책 20권"),
    ("기술 스택", "Next.js 16, Claude Haiku 4.5, 한국관광공사 TourAPI, Python 자동화"),
]
for i, (k, v) in enumerate(rows):
    t2.rows[i].cells[0].text = k
    t2.rows[i].cells[1].text = v

# ─── 2. 창업 동기 ───
h1("2. 창업 동기 및 지역적 배경")
para("신청인은 경주 본사를 둔 1인 디지털 스튜디오 '쿤스튜디오'의 대표이다.", bold=True)
para("경주는 연간 1,500만 명 이상의 관광객이 방문하는 세계문화유산 도시지만, 실제 방문객들은 다음과 같은 구조적 불편을 겪고 있다:")
bullet("경주 문화재·관광지 정보가 공공 API(한국관광공사 TourAPI)에 흩어져 있어 일반 사용자 접근이 어려움")
bullet("외국인 방문객을 위한 한영 이중 큐레이션 부재 (경주는 한류·K-드라마 주요 촬영지)")
bullet("가족·솔로·커플·외국인 등 방문자 유형별 맞춤 경로 추천 서비스 없음")
bullet("전통문화(사주·명리) 체험이 관광 상품으로 충분히 상품화되지 않음")
para("본 창업은 경주의 1,500년 신라 역사·문화 자산을 AI 큐레이션으로 재구성하여 지역경제 활성화와 관광객 만족도 제고를 동시에 달성한다.")

# ─── 3. 핵심 아이템 ───
h1("3. 핵심 아이템 — KORLENS 경주 특화 + 천명당")

h2("3-1. 메인: KORLENS (korlens.vercel.app)")
para("KORLENS는 한국관광공사 TourAPI 4종과 Claude Haiku 4.5 AI를 결합한 '4관점 AI 관광 큐레이션' 서비스다.", bold=True)
bullet("17개 광역시도 중 **경주 심층 특화 확장**이 본 사업의 핵심 목표")
bullet("4가지 관점 매트릭스: 외국인·커플·가족·솔로 — 동일 장소라도 관점별 다른 하이라이트")
bullet("AI 768개 경주 관광 하이라이트 사전 생성 (Claude Haiku 캐싱, 응답속도 즉시)")
bullet("한·영 이중 언어 토글 + 카카오·네이버 지도 딥링크")
bullet("혼잡도 3단계 필터 + 열린관광 API 연동 (무장애 편의시설 태그)")
bullet("AI 큐레이터 챗봇 (장소 카드 인라인 렌더)")
bullet("**이미 2026 한국관광공사 관광데이터 활용 공모전 제출 완료** (2026-04-18)")

h2("3-2. 서브: 천명당 (K-전통문화 체험)")
para("천명당은 신라 왕조 1,500년 전통을 현대적으로 재해석한 AI 기반 사주·관상·손금 서비스이다.", bold=True)
bullet("방문객 대상 '경주 역사지 + 전통 운세' 체험 패키지 기획")
bullet("불국사·석굴암·첨성대 등 관광 동선 연계")
bullet("외국인 대상 '한국 전통 문화 체험' 스토리텔링 (영어 지원)")

h2("3-3. 두 서비스의 지역 융합 시너지")
para("KORLENS(장소 큐레이션) + 천명당(전통 문화 체험) = 경주 방문 일정 전체를 커버하는 종합 플랫폼.", bold=True)
bullet("아침 관광 추천 (KORLENS) → 점심 맛집 → 오후 사주·관상 체험 (천명당) → 저녁 포토스팟")
bullet("외국인 패키지: 영문 관광 설명 + 한국 전통 운세 경험")

# ─── 4. 지역성 분석 (핵심) ───
h1("4. 지역성 분석 — 경주 특화 전략 (배점 50점 핵심)")

h2("4-1. 경주 자원·문화 이해도")
para("경주는 국내 유일 'UNESCO 세계문화유산 2건 보유' 지역이다.")
t3 = doc.add_table(rows=5, cols=2); t3.style = "Light Grid"
rows = [
    ("세계문화유산", "불국사·석굴암 (1995 등재), 경주역사지구 (2000 등재)"),
    ("지정문화재", "2,076건 (국가·도·시 지정 포함, 경주시 공식)"),
    ("연간 관광객", "연 1,500만 명+ (문화체육관광부 통계)"),
    ("한류 촬영지", "'더 킹: 영원의 군주', '킹덤' 등 다수"),
    ("특산물·문화", "황남빵·경주법주·신라토기, 전통 사주·명리 문화"),
]
for i, (k, v) in enumerate(rows):
    t3.rows[i].cells[0].text = k
    t3.rows[i].cells[1].text = v

h2("4-2. 지역 자원 활용 방안")
bullet("한국관광공사 TourAPI + 경주시 문화관광홈페이지 데이터 통합")
bullet("경주시립박물관·국립경주박물관 유물 정보 연계 (AI 해설)")
bullet("불국사·석굴암·첨성대·대릉원 등 주요 관광지 4관점 큐레이션")
bullet("경주 특산물 (황남빵·경주법주) 매장 자동 추천 (동선 연계)")
bullet("경주 전통 사주·명리 문화를 천명당으로 관광상품화")

h2("4-3. 지역경제 활성화 기여도")
para("본 사업을 통해 예상되는 지역경제 효과:", bold=True)
bullet("외국인 관광객 체류시간 평균 1.2일 → 2일+ 확대 (AI 추천 동선으로)")
bullet("관광객 1인당 지출액 증가 (소규모 로컬 상점 노출 증대)")
bullet("2026년 목표: 월간 방문자 10,000명 → 경주 로컬 상점 노출 약 30만 건/월")
bullet("지역 크리에이터 2명 이상 외주 협업 (영상·콘텐츠)")
bullet("경주 소상공인 대상 KORLENS 장소 등록 무료 (경주 상공회의소 제휴 목표)")

# ─── 5. 사업 추진 전략 ───
h1("5. 사업 추진 전략 (2026.5 ~ 12)")

h2("5-1. 월별 로드맵")
t4 = doc.add_table(rows=9, cols=3); t4.style = "Light Grid"
hdr = t4.rows[0].cells
hdr[0].text = "시기"; hdr[1].text = "주요 활동"; hdr[2].text = "목표 지표"
data = [
    ("2026-05", "KORLENS 경주 특화 확장 기획 / 경주시·관광공사 제휴 제안", "파트너 2곳 접촉"),
    ("2026-06", "경주 768개 장소 심층 데이터 수집 + Claude Haiku 파인튜닝", "장소 DB 2배 확장"),
    ("2026-07", "천명당 '경주 전통 운세' 테마 패키지 출시", "방문자 MAU 1,000"),
    ("2026-08", "여름 성수기 외국인 마케팅 (영어권 SNS)", "해외 트래픽 10%"),
    ("2026-09", "무장애 편의시설 DB 완성 + 고령·장애인 타겟 확장", "접근성 태그 100%"),
    ("2026-10", "공모전 1차·최종심사 대비 + KPI 기반 개선", "MAU 5,000"),
    ("2026-11", "경주 축제 연계 특별 큐레이션 (벚꽃·단풍)", "방문자 MAU 10,000"),
    ("2026-12", "연말 지역 상점 종합 리포트 + B2B 세일즈", "로컬 제휴 20곳"),
]
for i, r_ in enumerate(data, 1):
    for j, v in enumerate(r_):
        t4.rows[i].cells[j].text = v

h2("5-2. 파트너사 협업 계획")
bullet("창작자 파트너: 경주 지역 영상·사진 크리에이터 2명 (외주)")
bullet("기술 파트너: Claude Agent SDK 기반 AI 서비스 확장")
bullet("지역 파트너 목표: 경주시 문화관광과, 경상북도 관광공사, 경주시 상공회의소")
bullet("숙박·레스토랑 지역 소상공인 10~20곳 직접 제휴 (서비스 상호 노출)")

# ─── 6. 마케팅 계획 ───
h1("6. 마케팅 계획 — 외국인·국내 가족층 이중 타겟")
bullet("외국인: Instagram(영어 콘텐츠), YouTube Shorts, Reddit r/Korea 등 해외 커뮤니티")
bullet("국내 가족층: 네이버 블로그, Threads, KakaoChannel 자동 발행 (Postiz 셀프호스트 가동 중)")
bullet("경주 호텔·게스트하우스 QR 코드 배치 (1차 20곳)")
bullet("한국관광공사 공모전 수상 시 공식 홍보 채널 활용")
bullet("천명당·KORLENS 크로스 프로모션 (사주앱 사용자 → KORLENS 경주 방문 유도)")

# ─── 7. 자금 집행 계획 ───
h1("7. 사업화 자금 집행 계획 (최대 5,000만원)")
para("본 사업 로컬기업 육성 사업의 국비 지원비율 80%에 따라 자부담 20%(총사업비 대비) 포함 계획:", bold=True)

t5 = doc.add_table(rows=9, cols=3); t5.style = "Light Grid"
hdr = t5.rows[0].cells
hdr[0].text = "항목"; hdr[1].text = "금액(만원)"; hdr[2].text = "상세 내용"
data = [
    ("① 개발·기술 고도화", "1,500", "KORLENS 경주 지역 DB 확장, AI 큐레이션 파인튜닝, 영문 고도화"),
    ("② 지역 데이터 수집", "600", "경주 768개+ 관광지 심층 취재 비용, 사진·영상 자료"),
    ("③ 외주 인력", "800", "지역 크리에이터 2명, 번역(영·중·베트남), 촬영 협조"),
    ("④ 마케팅·홍보", "1,200", "해외 SNS 광고, Threads/인스타 유료 홍보, 호텔 QR 배포"),
    ("⑤ 법률·인증", "200", "관광 사업 컨설팅, 개인정보 관련 법무 검토"),
    ("⑥ 인프라", "500", "Vercel Pro, Claude API 크레딧, 도메인/SSL"),
    ("⑦ 사업비 지급보증보험증권 수수료", "100", "e-나라도움 지급 보증 수수료"),
    ("⑧ 예비비", "100", ""),
]
for i, r_ in enumerate(data, 1):
    for j, v in enumerate(r_):
        t5.rows[i].cells[j].text = v
para("합계: 5,000만원 (국비 4,000만 + 자부담 1,000만, 현금·현물 혼합 가능)")

# ─── 8. 성장 목표 ───
h1("8. 성장 목표 및 기대효과")

h2("8-1. 정량 목표 (2026년 12월)")
bullet("MAU 10,000명 (경주 방문 관광객 중 5% 도달)")
bullet("외국인 사용자 비중 20% (2,000명)")
bullet("B2B 제휴 경주 소상공인 20곳 이상")
bullet("KORLENS 평균 체류시간 5분+ (관광정보 앱 평균 2.3분 대비)")

h2("8-2. 정성·사회적 효과")
bullet("경주 관광 데이터 공공성 강화 (무장애·다언어 접근성)")
bullet("외국인 관광객의 경주 체류 시간 평균 1.2일 → 2일 이상")
bullet("경주 소상공인 온라인 노출 확대 (문화재청·지자체 한계 보완)")

h2("8-3. 확장·지속 가능성")
bullet("공모전 수상 시 한국관광공사 공식 파트너십 가능")
bullet("2027년: 17개 광역시도 확장 (제주·강원 등 타 지역 모델화)")
bullet("2028년: 해외 한국 관광 플랫폼 확장 (일본·베트남 진출)")

# ─── 9. 팀 구성 ───
h1("9. 팀 구성 및 파트너")

h2("9-1. 대표자")
bullet("홍덕훈 (1985생) — AI 자동화 + Python + Next.js 실전 역량")
bullet("KORLENS 단독 개발·배포 (Vercel 프로덕션 운영 중)")
bullet("Claude Agent SDK 기반 13개 부서 자동화 인프라 구축")

h2("9-2. 파트너 네트워크")
bullet("한국관광공사 TourAPI (공식 연동 중)")
bullet("경주시·경상북도 (공모전 매칭 공식 채널)")
bullet("지역 크리에이터 (외주 협업 예정)")

# ─── 10. 리스크 및 대응 ───
h1("10. 리스크 분석 및 대응 전략")

t6 = doc.add_table(rows=6, cols=3); t6.style = "Light Grid"
hdr = t6.rows[0].cells
hdr[0].text = "리스크"; hdr[1].text = "영향"; hdr[2].text = "대응"
data = [
    ("TourAPI 정책 변경", "중", "자체 장소 DB 병행 수집, 다양한 공공 API 병행"),
    ("관광객 시즌 변동", "중", "계절별 큐레이션 (축제·벚꽃·단풍) + 비수기 로컬 콘텐츠"),
    ("1인 운영 업무 과부하", "상", "외주 크리에이터 협업 + AI 자동화 확대"),
    ("경쟁 서비스 (트리플·마이리얼트립)", "하", "경주 심층 특화로 차별화 (전국형 대비)"),
    ("외국인 방문 감소 (대외 요인)", "중", "국내 가족·커플 세그먼트 병행 개발"),
]
for i, r_ in enumerate(data, 1):
    for j, v in enumerate(r_):
        t6.rows[i].cells[j].text = v

# ─── 11. 부록 ───
h1("11. 부록 — 실적·채널")
bullet("KORLENS 운영 중: https://korlens.vercel.app")
bullet("2026 한국관광공사 공모전 접수 완료: 2026-04-18")
bullet("천명당 웹: https://ghdejr11-beep.github.io/cheonmyeongdang/")
bullet("경주시 문화관광 데이터: https://www.gyeongju.go.kr/tour/")
bullet("한국관광공사 TourAPI: https://api.visitkorea.or.kr/")
bullet("KDP 출판 20권, Gumroad 디지털 상품 8종 (창업 이전 실적)")

out = Path(__file__).parent / "사업계획서_로컬유형_KORLENS_v1.docx"
doc.save(out)
print(f"저장: {out}")
print(f"크기: {out.stat().st_size/1024:.0f} KB")
