"""
2026 강한소상공인 혁신형 사업계획서 초안 생성
"""
import sys
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

doc = Document()

style = doc.styles["Normal"]
style.font.name = "맑은 고딕"
style.font.size = Pt(10)


def h1(text):
    p = doc.add_heading(text, level=1)
    if p.runs:
        p.runs[0].font.name = "맑은 고딕"
    return p


def h2(text):
    p = doc.add_heading(text, level=2)
    if p.runs:
        p.runs[0].font.name = "맑은 고딕"
    return p


def para(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "맑은 고딕"
    r.font.size = Pt(10)
    if bold:
        r.bold = True
    return p


def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    for r in p.runs:
        r.font.name = "맑은 고딕"
        r.font.size = Pt(10)
    return p


# 표지
title = doc.add_heading("", 0)
run = title.add_run("2026년 소상공인 도약 지원 사업\n(강한소상공인 성장지원 — 혁신 유형)\n사업계획서")
run.font.name = "맑은 고딕"
run.font.size = Pt(18)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
doc.add_paragraph()

info_table = doc.add_table(rows=5, cols=2)
info_table.style = "Light Grid"
data = [
    ("상호", "쿤스튜디오"),
    ("대표자", "홍덕훈 (1985-02-03)"),
    ("사업자등록번호", "552-59-00848"),
    ("사업장", "경상북도 경주시 외동읍 제내못안길 26-62"),
    ("신청 유형", "강한소상공인 성장지원 — 혁신 유형"),
]
for i, (k, v) in enumerate(data):
    info_table.rows[i].cells[0].text = k
    info_table.rows[i].cells[1].text = v

doc.add_page_break()

# 1. 기업 일반 현황
h1("1. 기업 일반 현황")
t = doc.add_table(rows=7, cols=2)
t.style = "Light Grid"
rows = [
    ("업태 / 종목", "정보통신업 / 응용 소프트웨어 개발 및 공급업"),
    ("개업일자", "2026-04-01 (창업 20일차)"),
    ("과세 유형", "간이과세자"),
    ("상시 근로자 수", "1명 (대표자)"),
    ("주요 채널", "Gumroad / Google Play / Amazon KDP / Threads / Telegram / KakaoChannel"),
    ("보유 서비스", "7개 (천명당, 세금N혜택, 보험다보여, KORLENS, HexDrop, 크티 프롬프트팩, KDP 전자책 20권)"),
    ("핵심 기술", "Claude Agent SDK, Python 자동화, AI 프롬프트 엔지니어링, CODEF API 연동"),
]
for i, (k, v) in enumerate(rows):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v

# 2. 창업 동기 및 배경
h1("2. 창업 동기 및 배경")
para("신청인은 2026년 4월 1일 경상북도 경주에 1인 AI 스튜디오 '쿤스튜디오'를 창업하였다.")
para("기존 예비창업 단계에서 다음과 같은 시장 불편과 기회를 발견하였다:", bold=True)
bullet("세무 서비스: 소상공인·프리랜서의 종합소득세 환급 누락 (삼쩜삼은 수수료 높음, 직접 신고는 난이도 높음)")
bullet("사주·관상 서비스: 앱 평균 월 4,000원 구독 구조로 1인 운영 한계")
bullet("보험 비교: 내국인 대상 플랫폼(보맵·굿리치)은 있지만 외국인(결혼이민·영주 188K+221K)을 위한 한국어·외국어 통합 비교 플랫폼은 공백")
bullet("관광 데이터: 지역 관광 추천이 분절되어 있어 통합 경험 제공 공백")
para("본 창업은 위 각 도메인에서 'AI + 자동화 + 소상공인'의 결합으로 기존 대기업이 놓친 틈새 시장을 겨냥한다.")

# 3. 사업 아이템 개요
h1("3. 사업 아이템 개요 및 혁신성")

h2("3-1. 핵심 아이템 — AI 세무·금융 자동화 플랫폼 (세금N혜택)")
para("쿤스튜디오의 대표 아이템은 AI 기반 소상공인 세무 자동화 서비스 '세금N혜택'이다.", bold=True)
bullet("CODEF API 연동으로 국세청 홈택스 실시간 조회 (환급금·세금신고·연말정산간소화)")
bullet("AI 챗봇(Claude Agent SDK)으로 세무 상담 자동화 (24시간 응답)")
bullet("9.9% 합리적 수수료 + 환급 0원 시 100% 환불 보장 (경쟁사 삼쩜삼 대비 혁신적 신뢰 보장)")
bullet("1단계: 개인 소상공인 환급 → 2단계: 외국인 전용 트랙 → 3단계: B2B 세무 API")

h2("3-2. 연계 아이템 — 도메인별 AI 서비스 6종")
para("본 창업은 단일 서비스가 아닌 '1인 AI 스튜디오'로, 세금N혜택을 주축으로 다음 도메인에 동시 진출한다.")
bullet("천명당 (AI 사주·관상·손금) — 앱인토스 미니앱 빌드 완료, 토스 5,100만 사용자 노출 채널 확보")
bullet("보험다보여 — 외국인 전용 트랙 기획 (F-6 188K·F-5 221K 세그먼트)")
bullet("KORLENS — 2026 한국관광공사 공모전 출품 완료 (4/18)")
bullet("KDP 전자책 20권 + Gumroad 디지털 상품 8종 + 크티 프롬프트팩 9종")
bullet("HexDrop 퍼즐 게임 + Faceless YouTube 자동 제작 파이프라인")

h2("3-3. 기술적 혁신성")
para("다음 기술 스택으로 '1인 운영 × 7개 서비스'라는 업계 이례적 구조를 달성하였다.", bold=True)
bullet("Claude Agent SDK 기반 부서별 AI 에이전트 (13개 부서 자동화)")
bullet("Python 자동화 파이프라인 (ceo-briefing·수익집계·경쟁사 인텔리전스·KDP 업로드)")
bullet("Postiz 셀프호스트로 Threads·Instagram·YouTube·Telegram·KakaoChannel 다채널 동시 발행")
bullet("Gumroad·쿠팡·Google Play Reporting API 통합으로 일일 수익 자동 집계")
bullet("CODEF 연동 세무·보험·은행 데이터 실시간 처리 아키텍처")

# 4. 시장 분석
h1("4. 시장 분석 및 목표 고객")

h2("4-1. 타겟 시장")
t = doc.add_table(rows=5, cols=3)
t.style = "Light Grid"
hdr = t.rows[0].cells
hdr[0].text = "서비스"
hdr[1].text = "타겟"
hdr[2].text = "시장 규모 (팩트)"
rows = [
    ("세금N혜택", "소상공인·프리랜서 환급 미신청자", "국세청 기준 연 2조원 환급 누락"),
    ("천명당", "20~40대 여성 사주·연애 관심자", "점신 앱인토스 입점 후 DAU 6배, 월 2.1억 매출 사례"),
    ("보험다보여 (외국인)", "F-6 결혼이민 + F-5 영주", "법무부 2025말 기준 188K + 221K = 40만 명"),
    ("KDP / Gumroad / 크티", "디지털 상품 구매층 (국내외)", "경쟁사 대비 한국어 AI 프롬프트팩 공백"),
]
for i, r in enumerate(rows, 1):
    for j, v in enumerate(r):
        t.rows[i].cells[j].text = v

h2("4-2. 경쟁사 분석 및 포지셔닝")
para("경쟁사 비교 결과 아래와 같이 포지셔닝한다.", bold=True)
bullet("세금N혜택: 삼쩜삼(수수료 15~30%) 대비 9.9% + 환불보장으로 신뢰 차별화")
bullet("천명당: 점신·포스텔러 대비 'B2C 앱 + API 기반 B2B 판매' 이중 수익 구조")
bullet("보험다보여: 보맵·굿리치 대비 '외국인 전용 라우트' 블루오션 선점")

# 5. 추진 전략
h1("5. 사업 추진 전략")

h2("5-1. 단계별 로드맵 (2026.5 ~ 2026.12)")
bullet("1단계 (5~7월): 세금N혜택 정식 출시 (CODEF 정식 승인 후), 1,000명 사전예약 완료")
bullet("2단계 (8~9월): 천명당 앱인토스 미니앱 출시 (토스 5,100만 사용자 노출), MAU 10,000 달성")
bullet("3단계 (10~11월): 외국인 보험다보여 MVP 출시, 다문화가족지원센터 제휴 5곳")
bullet("4단계 (12월): B2B 세무 API 2곳 계약 체결, 월 매출 2,000만원 돌파 (ARR 2.4억 기업화)")

h2("5-2. 파트너 협업 계획")
para("본 사업의 '혁신 유형' 권장 요건인 파트너사 협업은 다음과 같이 구성한다.", bold=True)
bullet("창작자 파트너: 콘텐츠 크리에이터 2명 (천명당 사주 콘텐츠 + 세무 영상)")
bullet("기술 파트너: CODEF API(헥토데이터) — 이미 정식 신청 완료")
bullet("지역 파트너: 경주시 소재 다문화가족지원센터·관광공사 (KORLENS 연계)")

# 6. 마케팅 계획
h1("6. 마케팅 계획")
bullet("Threads·Instagram 다채널 자동 발행 (Postiz 이미 가동 중)")
bullet("쿠팡 파트너스 + 크티 B2B 프롬프트팩 교차 판매")
bullet("앱인토스 5,100만 토스 사용자 노출 (천명당 미니앱)")
bullet("KDP 20권 역으로 세금N혜택·천명당 홍보 채널 활용")
bullet("네이버 블로그·유튜브 faceless 자동 발행 (lyrics_watcher.py)")

# 7. 사업화 자금
h1("7. 사업화 자금 집행 계획 (최대 1억원)")
t = doc.add_table(rows=7, cols=3)
t.style = "Light Grid"
hdr = t.rows[0].cells
hdr[0].text = "항목"
hdr[1].text = "금액 (만원)"
hdr[2].text = "내용"
rows = [
    ("① 개발·기술 고도화", "3,500", "세금N혜택 UI/UX 리디자인, 외국인 보험 라우트 개발, B2B API 서버 구축"),
    ("② 마케팅·홍보", "2,500", "Meta/인스타/YouTube 유료 광고, 앱인토스 프로모션, 창작자 파트너십"),
    ("③ 인건비 (외주)", "2,000", "영상 제작자 1명, 디자이너 1명, 외국어 번역 (영/중/베트남)"),
    ("④ 법률·회계", "500", "서비스 약관, 외국인 보험 규제 컨설팅, 세무 검토"),
    ("⑤ 인증·인프라", "1,000", "CODEF 정식 구독, Postiz·AWS·Vercel 프로덕션 이전, 도메인·SSL"),
    ("⑥ 자부담금", "0", "강한소상공인(혁신 유형) 국비 100% 지원 해당"),
]
for i, r in enumerate(rows, 1):
    for j, v in enumerate(r):
        t.rows[i].cells[j].text = v
para("(합계 1억원 기준, 실제 배정액에 따라 우선순위 조정)")

# 8. 성장 목표
h1("8. 성장 목표 및 기대 효과")

h2("8-1. 재무 목표 (2026년 말)")
bullet("월 매출 (MRR) 2,000만원 돌파 → ARR 2.4억 기업 진입")
bullet("누적 유료 사용자 5,000명 (세금N혜택 1,500 + 천명당 3,000 + 기타 500)")
bullet("B2B 세무 API 2건 계약 (각 건당 월 200만원 × 2)")

h2("8-2. 사회·지역적 기대 효과")
bullet("경북 경주 본사 유지 → 지역 일자리 최소 2명 (외주 인력) 간접 창출")
bullet("외국인(결혼이민·영주) 40만 명 소외 계층 보험 접근성 개선")
bullet("소상공인 세무 환급 누락 2조원 시장 중 유의미한 환급 회수 기여")

h2("8-3. 확장·지속 가능성")
bullet("2027년 TIPS 일반 트랙 도전 (매출 실적·기술 검증 기반)")
bullet("2028년 해외 확장 (일본·베트남 — KORLENS + 천명당 이미 다국어 로드맵)")
bullet("7개 서비스 서로 크로스셀 구조로 단일 실패 리스크 최소화")

# 9. 팀 구성
h1("9. 팀 구성 및 역량")

h2("9-1. 대표자 경력")
bullet("홍덕훈 (1985생) — AI 자동화·Python 개발 숙련")
bullet("Claude Agent SDK 기반 멀티 에이전트 아키텍처 설계·운영")
bullet("2026년 4월 창업과 동시에 13개 부서 자동화 인프라 구축 완료")
bullet("이미 Gumroad 8상품, Google Play 2앱(내부테스트), KDP 20권 출판 실적")

h2("9-2. 외부 전문가 네트워크")
bullet("세무사 자문 (세금N혜택 법적 리뷰)")
bullet("쿠팡 파트너스 마케팅 네트워크")
bullet("한국관광공사 공모전 네트워크 (KORLENS)")

# 10. 부록
h1("10. 부록 — 증빙 및 채널")
bullet("세금N혜택 랜딩: https://ghdejr11-beep.github.io/cheonmyeongdang/departments/tax/src/")
bullet("천명당 웹: https://ghdejr11-beep.github.io/cheonmyeongdang/")
bullet("세금N혜택 API (Vercel): https://tax-n-benefit-api.vercel.app")
bullet("Gumroad 상점: 등록된 상품 8종 (AI 부업 시스템 · 노션 템플릿 등)")
bullet("KDP 출판 도서: 20권 (저자명 Deokgu Studio / Deokhun Hong)")
bullet("KORLENS 공모전 제출 확인: 2026-04-18")

out = Path(__file__).parent / "사업계획서_초안_v1.docx"
doc.save(out)
print(f"저장: {out}")
print(f"크기: {out.stat().st_size/1024:.0f} KB")
