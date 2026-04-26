"""
2026 소상공인 도약 지원사업 — 강한소상공인 성장지원 (혁신 유형) v2
v1 대비 경쟁사 팩트 강화 + 재무 상세 + 리스크 + 혁신 정의 (라이프스타일) 맞춤
"""
import sys
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

doc = Document()
style = doc.styles["Normal"]
style.font.name = "맑은 고딕"
style.font.size = Pt(10)


def h1(text):
    p = doc.add_heading(text, level=1)
    if p.runs: p.runs[0].font.name = "맑은 고딕"
    return p


def h2(text):
    p = doc.add_heading(text, level=2)
    if p.runs: p.runs[0].font.name = "맑은 고딕"
    return p


def para(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "맑은 고딕"; r.font.size = Pt(10)
    if bold: r.bold = True
    return p


def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    for r in p.runs:
        r.font.name = "맑은 고딕"; r.font.size = Pt(10)
    return p


# 표지
t = doc.add_heading("", 0)
r = t.add_run("2026년 소상공인 도약 지원 사업\n(강한소상공인 성장지원 — 혁신 유형)\n사업계획서 v2")
r.font.name = "맑은 고딕"; r.font.size = Pt(18)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
doc.add_paragraph()

info = doc.add_table(rows=6, cols=2); info.style = "Light Grid"
rows = [
    ("상호", "쿤스튜디오"),
    ("대표자", "홍덕훈 (1985-02-03)"),
    ("사업자등록번호", "552-59-00848"),
    ("사업장", "경상북도 경주시 외동읍 제내못안길 26-62"),
    ("핵심 아이템", "AI 기반 1인 창업·프리랜서 라이프스타일 콘텐츠 플랫폼"),
    ("신청 유형", "강한소상공인 혁신 유형 (국비 100% / 최대 1억)"),
]
for i, (k, v) in enumerate(rows):
    info.rows[i].cells[0].text = k
    info.rows[i].cells[1].text = v
doc.add_page_break()

# 1. 기업 일반 현황
h1("1. 기업 일반 현황")
t2 = doc.add_table(rows=8, cols=2); t2.style = "Light Grid"
rows = [
    ("업태·종목", "정보통신업 / 응용 소프트웨어 개발 및 공급업"),
    ("개업일자", "2026-04-01 (창업 20일차)"),
    ("과세유형", "간이과세자 (소상공인 확정)"),
    ("상시근로자", "1명 (대표자)"),
    ("주요 서비스", "세금N혜택 / 천명당 / 보험다보여 / KDP 20권 / Gumroad 8상품 / 크티 9팩"),
    ("서비스 본질", "AI로 재무·취향·건강·지식 등 개인 라이프스타일 차별화 콘텐츠"),
    ("기술 스택", "Claude Agent SDK, CODEF API, Python 자동화, 13개 부서 에이전트 인프라"),
    ("이미 확보 자산", "토스 인앱결제 허가, 카카오 비즈채널, Postiz SNS 자동화, Meta 비즈 포트폴리오"),
]
for i, (k, v) in enumerate(rows):
    t2.rows[i].cells[0].text = k
    t2.rows[i].cells[1].text = v

# 2. 창업 동기
h1("2. 창업 동기 — 개인 라이프스타일 차별화 필요성")

para("혁신 유형의 정의는 '개인의 라이프스타일(취향, 가치) 등을 반영하여 차별화된 콘텐츠로 비즈니스모델을 고도화시키는 유형'이다. 본 창업은 '1인 창업자·프리랜서·디지털 노마드'라는 신(新) 라이프스타일을 타겟한다.", bold=True)

h2("2-1. 타겟 라이프스타일 시장 규모 (팩트)")
bullet("국내 1인 창업자 증가: 1인 기업은 2020년 51만개 → 2024년 77만개 (중기부)")
bullet("프리랜서·N잡러 증가: 국세청 종합소득세 신고자 매년 10%+ 성장")
bullet("디지털 부업 관심층 (AI 시대): Google Trends 'ai side hustle' 주간 +380% 상승")
bullet("한국 점술·운세 시장 규모: **1조 4천억** (매거진한경 2024)")
bullet("세금 환급 시장 (삼쩜삼 기준): 누적 가입자 **2,300만명**, 매출 2023년 상반기 777억원")

h2("2-2. 기존 서비스의 한계 (시장 불편)")
bullet("삼쩜삼: 수수료 **16.5%** (한국경제 2025 공식), 환급 외 재무 컨텍스트 없음")
bullet("점신·포스텔러: 연매출 830억이지만 여성 20대 타겟 국한, 30~40대 창업자·프리랜서 대상 부재")
bullet("보맵·굿리치: 내국인 대상, 외국인(F-6 188K·F-5 221K) 대상 한국어·외국어 통합 서비스 부재")
bullet("KDP·Gumroad: 한국어 AI 프롬프트·템플릿 상품 공백 (영어권 포화 대비)")

# 3. 핵심 아이템
h1("3. 핵심 아이템 개요")

h2("3-1. 메인: 세금N혜택 (AI 세무·재무 통합 플랫폼)")
para("세금N혜택은 CODEF API(헥토데이터)와 Claude Agent SDK를 결합한 1인 창업자·프리랜서 전용 세무 자동화 서비스이다.", bold=True)

t3 = doc.add_table(rows=6, cols=3); t3.style = "Light Grid"
hdr = t3.rows[0].cells
hdr[0].text = "항목"; hdr[1].text = "삼쩜삼 (경쟁사)"; hdr[2].text = "세금N혜택 (쿤스튜디오)"
data = [
    ("수수료율", "16.5% (공식)", "**9.9%** (40% 저렴)"),
    ("환급 0원 시", "전액 환불", "전액 환불 + 자동 대안 제시"),
    ("AI 상담", "제한적 FAQ", "Claude Agent SDK **24시간 맞춤 챗봇**"),
    ("외국인 대응", "없음", "영문·중문·베트남어 트랙 (2단계)"),
    ("타겟", "전 국민", "**1인 창업자·프리랜서** 라이프스타일 특화"),
]
for i, r_ in enumerate(data, 1):
    for j, v in enumerate(r_):
        t3.rows[i].cells[j].text = v

h2("3-2. 서브: 천명당 (K-문화 라이프스타일 콘텐츠)")
bullet("AI 기반 사주·관상·손금 서비스 (Play Store 비공개 테스트 중)")
bullet("앱인토스 미니앱 빌드 완료 — 토스 5,100만 사용자 노출 채널 (앱인토스 미니앱 중 75%가 실매출 발생)")
bullet("비교: 점신(테크랩스) 연매출 830억 / 누적 1,500만 다운로드 — 경쟁 가능한 틈새")

h2("3-3. 서브: 보험다보여 (외국인·소외계층 라이프스타일)")
bullet("F-6 결혼이민 188,105명 + F-5 영주 221,021명 (법무부 2025.12 기준)")
bullet("보맵·굿리치 등 한국어 전용 서비스 공백을 영문·중문·베트남어로 채움")
bullet("다문화가족지원센터·산부인과 제휴 목표")

h2("3-4. 기타: KDP 20권 + Gumroad 8상품 + 크티 9팩")
bullet("창업 전 준비 기간 실적물 — 제품화 시스템 검증됨")
bullet("AI 부업·프롬프트·노션 템플릿 카테고리 집중 (라이프스타일 디지털 상품)")

# 4. 시장 분석
h1("4. 시장 분석 및 경쟁력")

h2("4-1. TAM / SAM / SOM (팩트 기반)")
t4 = doc.add_table(rows=5, cols=4); t4.style = "Light Grid"
hdr = t4.rows[0].cells
hdr[0].text = "구분"; hdr[1].text = "규모"; hdr[2].text = "산출 근거"; hdr[3].text = "쿤스튜디오 전략"
data = [
    ("TAM", "약 2조원", "세금 환급 + 운세 + 외국인 보험 통합", "3대 도메인 교차"),
    ("SAM", "약 3,000억", "30~45세 1인 창업자·프리랜서 시장", "세금N혜택 메인"),
    ("SOM", "약 30억", "SAM의 1% (3년 목표)", "MAU 5만 × ARPU 5만"),
    ("2026년 목표", "월 2,000만원", "MRR 2,000만 × 12 = 2.4억 ARR", "4대 채널 다각화"),
]
for i, r_ in enumerate(data, 1):
    for j, v in enumerate(r_):
        t4.rows[i].cells[j].text = v

h2("4-2. 차별화 포인트 — 혁신성 (배점 15점)")
para("① 수수료 혁신: 삼쩜삼 16.5% vs 당사 9.9% → **40% 저렴**", bold=True)
para("② AI 챗봇: 공정거래위원회·국세청 FAQ 자동 파싱 + Claude Agent 24시간 상담", bold=True)
para("③ 환불보장 + 대안제시: 환급 0원 시 100% 환불 + '다른 혜택 자동 찾아주기'", bold=True)
para("④ 라이프스타일 밀착: 1인 창업자·프리랜서 전용 '수익화·세무·투자' 종합 플랫폼 (경쟁사 없음)", bold=True)

# 5. 추진 전략
h1("5. 사업 추진 전략 (2026.5 ~ 12)")

h2("5-1. 월별 로드맵")
t5 = doc.add_table(rows=9, cols=3); t5.style = "Light Grid"
hdr = t5.rows[0].cells
hdr[0].text = "시기"; hdr[1].text = "주요 활동"; hdr[2].text = "목표 지표"
data = [
    ("2026-05", "CODEF 정식 승인 + 세금N혜택 유료 베타 출시", "결제 파일럿 100건"),
    ("2026-06", "천명당 앱인토스 미니앱 공식 론칭 (토스 노출 시작)", "DAU 1,000"),
    ("2026-07", "크티 프롬프트팩 3종 신규 + Gumroad 다국어 (영/일) 확장", "매출 월 500만"),
    ("2026-08", "세금N혜택 외국인 트랙 MVP (영·중·베 번역)", "외국인 유저 비중 5%"),
    ("2026-09", "AI 챗봇 고도화 (Claude Agent SDK 2단계)", "CS 자동화 90%"),
    ("2026-10", "B2B 세무 API 1곳 계약 (중소기업 10곳 대상)", "월매출 1,500만"),
    ("2026-11", "외국인 보험다보여 MVP 출시 + 다문화 제휴 5곳", "가입자 누적 2,000"),
    ("2026-12", "월 MRR 2,000만 돌파 + TIPS·디딤돌 R&D 준비", "ARR 2.4억 진입"),
]
for i, r_ in enumerate(data, 1):
    for j, v in enumerate(r_):
        t5.rows[i].cells[j].text = v

h2("5-2. 파트너 협업 전략 (혁신 유형 권장 요건)")
t6 = doc.add_table(rows=4, cols=3); t6.style = "Light Grid"
hdr = t6.rows[0].cells
hdr[0].text = "파트너 유형"; hdr[1].text = "대상"; hdr[2].text = "역할·기여"
data = [
    ("창작자 (2명)", "영상 크리에이터 + 디자이너", "세금N혜택 콘텐츠 + 천명당 UX 고도화"),
    ("스타트업 (1곳)", "세무사·회계 에이전시", "세무 전문성 검증 + B2B 채널"),
    ("이업종 소상공인 (2곳)", "지역 카페·공방", "세금N혜택 파일럿 고객 + 컨텐츠 제휴"),
]
for i, r_ in enumerate(data, 1):
    for j, v in enumerate(r_):
        t6.rows[i].cells[j].text = v

# 6. 마케팅
h1("6. 마케팅·고객 획득 계획")

h2("6-1. 채널별 예산 및 CAC 목표")
t7 = doc.add_table(rows=6, cols=4); t7.style = "Light Grid"
hdr = t7.rows[0].cells
hdr[0].text = "채널"; hdr[1].text = "월 예산"; hdr[2].text = "목표 CAC"; hdr[3].text = "메시지"
data = [
    ("Meta(인스타) 유료", "50만", "8,000원", "'삼쩜삼보다 40% 싸게'"),
    ("Threads 오가닉", "0", "3,000원", "AI 부업·창업 콘텐츠"),
    ("YouTube Shorts", "30만", "10,000원", "세무 팁 숏폼"),
    ("앱인토스 노출", "0", "2,000원", "5,100만 토스 사용자"),
    ("쿠팡 파트너스 + 블로그", "0", "5,000원", "실제 재테크 상품 연계"),
]
for i, r_ in enumerate(data, 1):
    for j, v in enumerate(r_):
        t7.rows[i].cells[j].text = v

h2("6-2. LTV 모델")
bullet("세금N혜택 평균 환급액 26만원 × 수수료 9.9% = 권당 매출 약 **2.5만원**")
bullet("Gumroad AOV 평균 3만원 (재구매율 30%)")
bullet("천명당 월 구독 4,900원 × 평균 6개월 = 3만원")
bullet("총 LTV 목표 (3년): 10만원 이상 → CAC 대비 10배+")

# 7. 자금 집행
h1("7. 사업화 자금 집행 계획 (최대 1억원, 국비 100%)")

t8 = doc.add_table(rows=8, cols=3); t8.style = "Light Grid"
hdr = t8.rows[0].cells
hdr[0].text = "항목"; hdr[1].text = "금액(만원)"; hdr[2].text = "상세"
data = [
    ("① 기술 개발", "3,500", "세금N혜택 UI/UX, 외국인 트랙, B2B API, AI 챗봇 고도화"),
    ("② 마케팅·홍보", "2,500", "Meta·YouTube 유료 광고, 앱인토스 프로모션, 파트너십"),
    ("③ 외주 인건비", "2,000", "영상 1명, 디자이너 1명, 번역 (영/중/베)"),
    ("④ 인프라·API", "1,000", "CODEF 정식, Vercel·AWS, Claude API 크레딧 (월 50만 × 12)"),
    ("⑤ 법무·회계·인증", "500", "약관, 외국인 서비스 규제, 개인정보보호 인증"),
    ("⑥ 데이터·콘텐츠", "300", "외국인 다문화 센터 협업 데이터, 경쟁사 벤치마크"),
    ("⑦ 예비비", "200", ""),
]
for i, r_ in enumerate(data, 1):
    for j, v in enumerate(r_):
        t8.rows[i].cells[j].text = v
para("합계: 10,000만원 (국비 100% — 자부담 없음)")

# 8. 성장 목표
h1("8. 성장 목표 및 손익분기")

h2("8-1. 재무 목표 (2026년 말)")
bullet("MRR 2,000만원 → ARR 2.4억원")
bullet("누적 유료 사용자 5,000명 (세금 1,500 + 천명당 3,000 + 기타 500)")
bullet("손익분기점 예상: 2026년 10월 (창업 7개월차)")

h2("8-2. 주요 KPI 월 추적")
bullet("세금N혜택: 신청자 / 환급 성공률 / 평균 환급액")
bullet("천명당: DAU / 결제전환율 / 앱인토스 노출 CTR")
bullet("Gumroad: 월 매출 / 신규 상품 수 / 다국어 비중")
bullet("쿠팡 파트너스: 누적 수수료 / 15만원 승인 달성")

h2("8-3. 확장·지속 가능성")
bullet("2027년: TIPS 일반 트랙 도전 (매출·기술 실적 기반)")
bullet("2028년: 일본·베트남 해외 확장 (세금N혜택 외국인 트랙 역수출)")
bullet("7개 서비스 상호 크로스셀 구조 → 단일 실패 리스크 최소화")

# 9. 팀 구성
h1("9. 팀 구성 및 역량")

h2("9-1. 대표자 역량")
bullet("AI 자동화·Python·Next.js 풀스택 역량")
bullet("Claude Agent SDK 기반 13개 부서 멀티 에이전트 아키텍처 운영")
bullet("창업 20일차 동시 실적: Gumroad 8상품 / KDP 20권 / Google Play 2앱 / 앱인토스 빌드")
bullet("Postiz 셀프호스트 (Railway) + Meta 비즈 포트폴리오 자동화 인프라")

h2("9-2. 외부 전문가 네트워크")
bullet("세무사 자문 (세금N혜택 법적 검토)")
bullet("CODEF(헥토데이터) 기술 파트너 — 정식 신청 완료")
bullet("한국관광공사 TourAPI 공식 연동 (KORLENS 실적)")
bullet("외국인 커뮤니티 접촉 채널 (다문화가족지원센터 227곳)")

# 10. 리스크 및 대응
h1("10. 리스크 분석 및 대응")

t9 = doc.add_table(rows=7, cols=3); t9.style = "Light Grid"
hdr = t9.rows[0].cells
hdr[0].text = "리스크"; hdr[1].text = "영향"; hdr[2].text = "대응 전략"
data = [
    ("CODEF 정식 승인 지연", "상", "샌드박스 유지 + 대체 API(홈택스 Open) 이원화"),
    ("경쟁사 수수료 인하", "중", "AI 상담·외국인 트랙 등 부가가치 차별화"),
    ("1인 운영 업무 과부하", "상", "외주 2명 + Claude Agent SDK 자동화 확대"),
    ("국세청 정책 변경", "중", "법무사·세무사 자문 루틴 구축"),
    ("앱인토스 정책 변경", "하", "Play Store·App Store 이원 운영"),
    ("외국인 규제(비자별)", "중", "전문 변호사 자문 + 정부 공식 데이터만 사용"),
]
for i, r_ in enumerate(data, 1):
    for j, v in enumerate(r_):
        t9.rows[i].cells[j].text = v

# 11. 부록
h1("11. 부록 — 실적·채널")
bullet("세금N혜택 랜딩: https://ghdejr11-beep.github.io/cheonmyeongdang/departments/tax/src/")
bullet("세금N혜택 API (Vercel): https://tax-n-benefit-api.vercel.app")
bullet("천명당 웹: https://ghdejr11-beep.github.io/cheonmyeongdang/")
bullet("KORLENS: https://korlens.vercel.app (한국관광공사 공모전 제출)")
bullet("Gumroad 상점: ghdejr.gumroad.com (8상품 등록)")
bullet("KDP 20권 (Amazon 저자: Deokgu Studio / Deokhun Hong)")
bullet("CODEF 정식 신청: 2026-04-17")

out = Path(__file__).parent / "사업계획서_혁신유형_세금N혜택_v2.docx"
doc.save(out)
print(f"저장: {out}")
print(f"크기: {out.stat().st_size/1024:.0f} KB")
