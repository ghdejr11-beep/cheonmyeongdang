"""
docx → PDF 변환 (한글 맑은 고딕)
사용: python _docx_to_pdf.py <input.docx> <output.pdf>
"""
import sys, os
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER

sys.stdout.reconfigure(encoding="utf-8")

# Register Korean font
pdfmetrics.registerFont(TTFont("Malgun", "C:/Windows/Fonts/malgun.ttf"))
pdfmetrics.registerFont(TTFont("MalgunBold", "C:/Windows/Fonts/malgunbd.ttf"))


def convert(docx_path, pdf_path):
    src = Document(docx_path)

    pdf = SimpleDocTemplate(
        pdf_path,
        pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm,
    )

    styles = {
        "Title": ParagraphStyle("Title", fontName="MalgunBold", fontSize=18, alignment=TA_CENTER, spaceAfter=20, leading=24),
        "H1":    ParagraphStyle("H1",    fontName="MalgunBold", fontSize=14, spaceBefore=12, spaceAfter=8, leading=18),
        "H2":    ParagraphStyle("H2",    fontName="MalgunBold", fontSize=12, spaceBefore=10, spaceAfter=6, leading=16),
        "Body":  ParagraphStyle("Body",  fontName="Malgun",     fontSize=10, spaceAfter=4, leading=14),
        "Bullet":ParagraphStyle("Bullet",fontName="Malgun",     fontSize=10, leftIndent=14, spaceAfter=2, leading=13),
    }

    story = []

    # Iterate both paragraphs and tables in document order
    body = src.element.body
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]

        if tag == "p":
            # find paragraph
            p = next((p for p in src.paragraphs if p._p is child), None)
            if p is None: continue
            txt = p.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            if not txt.strip():
                story.append(Spacer(1, 6))
                continue
            sname = p.style.name
            if "Title" in sname:
                story.append(Paragraph(txt, styles["Title"]))
            elif "Heading 1" in sname:
                story.append(Paragraph(txt, styles["H1"]))
            elif "Heading 2" in sname or "Heading" in sname:
                story.append(Paragraph(txt, styles["H2"]))
            elif "Bullet" in sname or "List" in sname:
                story.append(Paragraph(f"• {txt}", styles["Bullet"]))
            else:
                story.append(Paragraph(txt, styles["Body"]))

        elif tag == "tbl":
            # find table
            tbl = next((t for t in src.tables if t._tbl is child), None)
            if tbl is None: continue
            data = []
            for row in tbl.rows:
                row_cells = []
                for cell in row.cells:
                    text = "\n".join(p.text for p in cell.paragraphs).strip()
                    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    row_cells.append(Paragraph(text or "&nbsp;", styles["Body"]))
                data.append(row_cells)
            if data:
                t = Table(data, hAlign="LEFT")
                t.setStyle(TableStyle([
                    ("GRID", (0,0), (-1,-1), 0.5, colors.grey),
                    ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
                    ("FONTNAME", (0,0), (-1,-1), "Malgun"),
                    ("FONTSIZE", (0,0), (-1,-1), 9),
                    ("VALIGN", (0,0), (-1,-1), "TOP"),
                    ("LEFTPADDING", (0,0), (-1,-1), 4),
                    ("RIGHTPADDING", (0,0), (-1,-1), 4),
                    ("TOPPADDING", (0,0), (-1,-1), 3),
                    ("BOTTOMPADDING", (0,0), (-1,-1), 3),
                ]))
                story.append(Spacer(1, 4))
                story.append(t)
                story.append(Spacer(1, 4))

    pdf.build(story)
    print(f"PDF created: {pdf_path}")
    print(f"  size: {os.path.getsize(pdf_path):,} bytes")


if __name__ == "__main__":
    if len(sys.argv) >= 3:
        convert(sys.argv[1], sys.argv[2])
    else:
        base = r"C:\Users\hdh02\Desktop\cheonmyeongdang\departments\tax\applications\gyeongbuk_doyak"
        convert(
            os.path.join(base, "사업계획서_혁신유형_세금N혜택_v2.docx"),
            os.path.join(base, "01_사업계획서_혁신_final.pdf"),
        )
