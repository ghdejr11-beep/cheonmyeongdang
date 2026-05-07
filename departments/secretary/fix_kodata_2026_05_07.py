"""KoDATA 빈양식 fix — 사과 메일 + 기업정보 채워진 DOCX 재발송 (2026-05-07).

문제: 5/5 발송 메일에 (양식)한국관광공사 기업정보 등록의뢰서.hwp (빈 양식) 첨부 → KoDATA 회신 "빈 양식 보냈다"

해결:
1. 기업정보가 모두 채워진 DOCX 신규 생성 (한글 표 형식)
2. 사업자등록증명 PDF 함께 재첨부
3. 사과 + 본문에 정보 평문 명시
4. 동일 수신자(find@kodata.co.kr) 재발송
"""
import os
import json
import base64
import mimetypes
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.table import WD_ALIGN_VERTICAL

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
TOKEN_PATH = os.path.join(SCRIPT_DIR, 'token.json')

FROM_EMAIL = 'ghdejr11@gmail.com'
FROM_NAME = '쿤스튜디오 홍덕훈'
TO_EMAIL = 'find@kodata.co.kr'
SUBJECT = '[재발송] 기업정보등록의뢰서 - 쿤스튜디오 (552-59-00848)'

DOWNLOADS = r'C:\Users\hdh02\Downloads'
OUT_DIR = os.path.join(SCRIPT_DIR, 'output')
os.makedirs(OUT_DIR, exist_ok=True)
DOCX_PATH = os.path.join(OUT_DIR, '기업정보등록의뢰서_쿤스튜디오_2026_05_07.docx')

# ====================== DOCX 생성 ======================
COMPANY = {
    '기업명 (국문)': '쿤스튜디오',
    '기업명 (영문)': 'KunStudio',
    '대표자 성명': '홍덕훈',
    '사업자등록번호': '552-59-00848',
    '개업일자': '2026-04-01',
    '사업장 주소': '(38204) 경상북도 경주시 외동읍 제내못안길 25-52',
    '업태': '정보통신업',
    '종목': '응용 소프트웨어 개발 및 공급업',
    '과세유형': '간이과세자',
    '사업 형태': '개인사업자 (1인 운영)',
    '대표 연락처': '010-4244-6992',
    '대표 이메일': 'ghdejr11@gmail.com',
    '주요 사업': '천명당(글로벌 사주 SaaS, 4언어), KORLENS(K-콘텐츠 큐레이션), 세금N혜택',
    '등록 사유': '한국관광공사 「2026 관광기업 데이터·AI 활용 지원 사업」 (마감 2026.05.20) 신청 자격 확보',
}


def build_docx():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(11)

    title = doc.add_heading('기업정보 등록의뢰서', level=1)
    sub = doc.add_paragraph()
    sub.add_run('한국평가데이터(KoDATA) 제출용 / 한국관광산업포털(투어라즈) 가입 사전등록').italic = True

    doc.add_paragraph()

    # 기업정보 표
    table = doc.add_table(rows=len(COMPANY) + 1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '항목'
    hdr[1].text = '내용'
    for c in hdr:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    for i, (k, v) in enumerate(COMPANY.items(), start=1):
        row = table.rows[i].cells
        row[0].text = k
        row[1].text = v

    doc.add_paragraph()
    doc.add_paragraph('※ 모든 정보는 사업자등록증 및 사업자등록증명서 기준 정확히 기재되었습니다.')
    doc.add_paragraph('※ 첨부: 사업자등록증명서.pdf (사실 확인용)')
    doc.add_paragraph()
    doc.add_paragraph(f'작성일: 2026-05-07')
    doc.add_paragraph(f'작성자: {COMPANY["대표자 성명"]} (대표)')

    doc.save(DOCX_PATH)
    print(f'  [DOCX] {DOCX_PATH} ({os.path.getsize(DOCX_PATH):,} bytes)')


# ====================== 메일 ======================
BODY = """안녕하십니까, KoDATA 담당자님.

쿤스튜디오 홍덕훈입니다.

먼저, 5월 5일 발송드린 메일에 빈 양식 파일이 첨부되어 큰 불편을 드린 점 진심으로 사과드립니다.
당시 한글(HWP) 편집이 어려워 본문에 정보를 기재하고 양식 파일은 참고용으로 첨부드린 것이었으나,
혼선을 드린 점 죄송합니다.

본 재발송 메일에는 기업정보를 모두 정확히 기재한 워드(DOCX) 문서와
사업자등록증명서(PDF)를 첨부하여 드립니다.
한글(HWP) 양식이 반드시 필요하시면 회신 부탁드립니다.
직접 작성하여 우편 또는 팩스로 발송 가능합니다.

▶ 등록 의뢰 정보 요약 (사업자등록증 기준)

  - 기업명         : 쿤스튜디오 (KunStudio)
  - 대표자         : 홍덕훈
  - 사업자등록번호 : 552-59-00848
  - 개업일자       : 2026-04-01
  - 사업장 주소    : (38204) 경상북도 경주시 외동읍 제내못안길 25-52
  - 업태           : 정보통신업
  - 종목           : 응용 소프트웨어 개발 및 공급업
  - 과세유형       : 간이과세자
  - 사업 형태      : 개인사업자 (1인 운영)
  - 휴대폰         : 010-4244-6992
  - 이메일         : ghdejr11@gmail.com
  - 주요 사업      : 천명당(글로벌 사주 SaaS), KORLENS, 세금N혜택

▶ 등록 사유
  한국관광공사 「2026 관광기업 데이터·AI 활용 지원 사업」 신청 자격 확보 (마감 2026.05.20 18:00)

▶ 첨부
  1. 기업정보등록의뢰서_쿤스튜디오_2026_05_07.docx  (모든 정보 기재 완료)
  2. 사업자등록증명서.pdf  (사실 확인)

빠른 등록 처리 부탁드리겠습니다. 거듭 양식 혼선 드린 점 사과드립니다. 감사합니다.

쿤스튜디오 홍덕훈
사업자번호: 552-59-00848
이메일: ghdejr11@gmail.com
휴대폰: 010-4244-6992
"""


def load_service():
    with open(TOKEN_PATH, 'r', encoding='utf-8') as f:
        data = json.load(f)
    creds = Credentials(
        token=data.get('token'),
        refresh_token=data.get('refresh_token'),
        token_uri=data.get('token_uri'),
        client_id=data.get('client_id'),
        client_secret=data.get('client_secret'),
        scopes=data.get('scopes'),
    )
    return build('gmail', 'v1', credentials=creds)


def build_message():
    msg = MIMEMultipart()
    msg['From'] = f'{FROM_NAME} <{FROM_EMAIL}>'
    msg['To'] = TO_EMAIL
    msg['Subject'] = SUBJECT
    msg.attach(MIMEText(BODY, 'plain', 'utf-8'))

    attachments = [
        DOCX_PATH,
        os.path.join(DOWNLOADS, '사업자등록증명_쿤스튜디오.pdf'),
    ]
    for path in attachments:
        if not os.path.isfile(path):
            print(f'  [WARN] 첨부 누락: {path}')
            continue
        ctype, _ = mimetypes.guess_type(path)
        if ctype is None:
            ctype = 'application/octet-stream'
        maintype, subtype = ctype.split('/', 1)
        with open(path, 'rb') as f:
            part = MIMEBase(maintype, subtype)
            part.set_payload(f.read())
        encoders.encode_base64(part)
        fname = os.path.basename(path)
        part.add_header('Content-Disposition', 'attachment', filename=('utf-8', '', fname))
        msg.attach(part)
        print(f'  [ATTACH] {fname} ({os.path.getsize(path):,} bytes)')
    return msg


def main():
    print('=== KoDATA 재발송 (빈양식 fix) ===')
    print('  Step 1: DOCX 생성')
    build_docx()
    print('  Step 2: 메일 빌드')
    print(f'  From: {FROM_NAME} <{FROM_EMAIL}>')
    print(f'  To  : {TO_EMAIL}')
    print(f'  Subj: {SUBJECT}')
    msg = build_message()
    raw = base64.urlsafe_b64encode(msg.as_bytes()).decode()
    print('  Step 3: 발송')
    service = load_service()
    res = service.users().messages().send(userId='me', body={'raw': raw}).execute()
    print(f'  [SENT] message_id={res.get("id")}')


if __name__ == '__main__':
    main()
